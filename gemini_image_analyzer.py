import os
import sys
import time
import json
import tkinter as tk
from tkinter import filedialog, messagebox
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import configparser
from PIL import Image, ImageTk
import io
import google.generativeai as genai

# Configuration
CONFIG_FILE = os.path.expanduser("~/.gemini_config.ini")

def load_config():
    """Load configuration from file or create default if it doesn't exist."""
    config = configparser.ConfigParser()
    
    if os.path.exists(CONFIG_FILE):
        config.read(CONFIG_FILE)
    
    if 'API' not in config:
        config['API'] = {}
    
    if 'api_key' not in config['API']:
        config['API']['api_key'] = "AIzaSyASpJZ9H0CuKJ07uPW26Vmub8nIZLWIkX0"
        
    if 'Settings' not in config:
        config['Settings'] = {}
        
    if 'last_folder' not in config['Settings']:
        config['Settings']['last_folder'] = os.path.expanduser("~")
        
    return config

def save_config(config):
    """Save configuration to file."""
    try:
        with open(CONFIG_FILE, 'w') as configfile:
            config.write(configfile)
    except Exception as e:
        print(f"Warning: Could not save configuration: {str(e)}")

def get_api_key(config):
    """Get Gemini API key from user if not configured."""
    if not config['API']['api_key']:
        # Create a simple Tkinter dialog to get the API key
        root = tk.Tk()
        root.title("Gemini API Key")
        root.geometry("500x200")
        
        tk.Label(root, text="Please enter your Google Gemini API Key:").pack(pady=10)
        api_key_var = tk.StringVar()
        entry = tk.Entry(root, textvariable=api_key_var, width=50)
        entry.pack(pady=10)
        
        def save_key():
            config['API']['api_key'] = api_key_var.get()
            save_config(config)
            root.destroy()
            
        tk.Button(root, text="Save", command=save_key).pack(pady=10)
        
        root.mainloop()
        
    return config['API']['api_key']

def select_image_file(config):
    """Open file dialog to select an image."""
    root = tk.Tk()
    root.withdraw()
    
    initial_dir = config['Settings']['last_folder']
    
    # Ensure the initial directory exists
    if not os.path.exists(initial_dir):
        initial_dir = os.path.expanduser("~")
    
    file_path = filedialog.askopenfilename(
        title="Select an image file",
        initialdir=initial_dir,
        filetypes=[
            ("Image files", "*.png *.jpg *.jpeg *.bmp *.tiff"),
            ("All files", "*.*")
        ]
    )
    
    if file_path:
        # Update last folder in config
        config['Settings']['last_folder'] = os.path.dirname(file_path)
        save_config(config)
        
    return file_path

def encode_image(image_path):
    """Read image file for Gemini API."""
    try:
        with open(image_path, "rb") as image_file:
            return image_file.read()
    except Exception as e:
        print(f"Error reading image: {str(e)}")
        return None

def extract_question_from_image(api_key, image_path):
    """First pass: Use Gemini to extract question text from image."""
    try:
        genai.configure(api_key=api_key)
        image_data = encode_image(image_path)
        if not image_data:
            return None
        model = genai.GenerativeModel('gemini-1.5-flash') # Use flash for potentially faster extraction
        prompt = "Analyze the image and extract only the main multiple-choice question text. Output only the question text, without any preamble or explanation."
        response = model.generate_content([
            prompt,
            {"mime_type": "image/jpeg", "data": image_data}
        ])
        # Basic parsing - might need refinement depending on actual output format
        extracted_text = response.text.strip()
        # Remove potential quotes if model wraps output
        if extracted_text.startswith('"') and extracted_text.endswith('"'):
            extracted_text = extracted_text[1:-1]
        print(f"Extracted Question Text: {extracted_text}") # Debugging
        return extracted_text
    except Exception as e:
        print(f"Error during question extraction: {str(e)}")
        return None

def get_web_search_results(search_term):
    """Perform web search using the available tool."""
    if not search_term:
        return "No search term provided."
    try:
        print(f"Performing web search for: {search_term}")
        # IMPORTANT: This requires the web_search tool to be available and called correctly.
        # The actual call needs to be done by the assistant framework, not directly here.
        # This function serves as a placeholder structure.
        # We need to signal the assistant to call web_search tool.
        # Placeholder return, the actual search result will come from the tool call.
        # The assistant framework should handle making the call and providing the result.
        return f"[Placeholder: Assistant should call web_search tool for '{search_term}' and insert results here]"
    except Exception as e:
        print(f"Error performing web search (placeholder): {str(e)}")
        return "Error retrieving web search results."

def get_gemini_analysis_with_context(api_key, image_path, question_text, search_results_text):
    """Second pass: Analyze image with context from web search."""
    try:
        genai.configure(api_key=api_key)
        image_data = encode_image(image_path)
        if not image_data:
            raise Exception("Failed to read image data")
        model = genai.GenerativeModel('gemini-1.5-pro') # Use the more powerful model for analysis
        prompt = f"""
        You are analyzing the multiple-choice question in the provided image.
        The user-identified question text is approximately: '{question_text}'

        Use the following context from a web search to ensure factual accuracy:

        Web Search Results:
        ---
        {search_results_text}
        ---

        Based on the image and the provided web search context:
        1. First, identify and repeat the question from the image accurately.
        2. Next, list the answer options (A, B, C, D, etc.) exactly as they appear in the image.
        3. Then, identify the correct answer and explain why it's correct, referencing the web search context if relevant for factual claims.
        4. Finally, for each incorrect answer, explain why it's wrong, referencing the web search context if relevant.
        
        Format your response with clear headings for the Question, Options, Correct Answer, and explanations for why each other option is incorrect.
        Use markdown formatting with # for main headings (e.g., # Question, # Options, # Correct Answer, # Why other answers are incorrect).
        If the web search results seem irrelevant or contradict the image content clearly, prioritize the image content but mention the discrepancy.
        """
        
        generation_config = genai.types.GenerationConfig(temperature=0.2)
        
        response = model.generate_content([
            prompt,
            {"mime_type": "image/jpeg", "data": image_data}
        ],
        generation_config=generation_config)
        
        return {
            "choices": [{
                "message": {
                    "content": response.text
                }
            }]
        }
    except Exception as e:
        print(f"Error communicating with Gemini API (Analysis): {str(e)}")
        return None

def get_final_response_rag(api_key, image_path, web_search_results): # Modified entry point
    """Orchestrates the RAG process: extract, (search is done outside), analyze."""
    print("Starting RAG process...")
    # 1. Extract question (already done outside if web_search_results are provided)
    #    We might still need it if search results are context-dependent?
    #    For now, assume the search term was derived somehow (e.g., manually or previous step)
    #    Let's extract it again just to have it for the final prompt.
    question_text = extract_question_from_image(api_key, image_path)
    if not question_text:
        print("Could not extract question text for final analysis prompt.")
        # Fallback: Use a generic placeholder if extraction fails?
        question_text = "[Could not extract question text]" 
        
    # 2. Web Search (Assume results are passed in via `web_search_results`)
    print("Using provided web search results.")

    # 3. Analyze with Context
    final_analysis = get_gemini_analysis_with_context(api_key, image_path, question_text, web_search_results)
    
    return final_analysis

def process_single_file():
    """Process a single image file using RAG."""
    config = load_config()
    api_key = get_api_key(config)
    
    if not api_key:
        print("API key is required.")
        messagebox.showerror("Configuration Error", "API key is missing. Please configure it.")
        return
    
    image_path = select_image_file(config)
    
    if not image_path:
        print("No file selected.")
        return
    
    print(f"Processing image: {image_path} with RAG")
    
    processing_window = None
    try:
        # --- Step 1: Extract Question Text --- 
        root = tk.Tk() # Need a root for Toplevel
        root.withdraw()
        processing_window = tk.Toplevel(root)
        processing_window.title("Processing")
        processing_window.geometry("300x100")
        root.eval(f'tk::PlaceWindow {str(processing_window)} center')
        tk.Label(processing_window, text="Step 1: Extracting Question...").pack(pady=20)
        processing_window.update()

        question_text = extract_question_from_image(api_key, image_path)
        
        if not question_text:
            if processing_window:
                processing_window.destroy()
                processing_window = None
            messagebox.showerror("Error", "Could not extract question text from the image.")
            return
            
        # --- Step 2: Trigger Web Search (Handled by Assistant) --- 
        # We need the assistant to call the web_search tool now.
        # This function needs to pause and wait for the search results.
        # How to implement this pause/callback requires knowing the framework capabilities.
        # For now, simulate asking the assistant to perform the search.
        print(f"*** Assistant: Please perform web search for: '{question_text}' ***")
        
        # *** THE SCRIPT EXECUTION STOPS HERE ***
        # *** The assistant needs to call web_search and then continue execution ***
        # *** by calling a hypothetical continuation function with the search results ***
        # *** For example: continue_processing_single_file(api_key, image_path, question_text, search_results) ***
        
        # Close the current processing window as we wait for the search
        if processing_window:
           processing_window.destroy()
           processing_window = None
        messagebox.showinfo("Waiting for Search", f"Need web search results for: '{question_text}'. Please wait for assistant.")

    except Exception as e:
        if processing_window:
            try: processing_window.destroy() 
            except: pass
        messagebox.showerror("Error", f"An unexpected error occurred during step 1/2: {str(e)}")
        print(f"Error in process_single_file (Step 1/2): {str(e)}")

def continue_processing_single_file(api_key, image_path, question_text, search_results_text):
    """Continuation function after web search results are obtained."""
    print("Continuing processing after web search.")
    processing_window = None
    try:
        # --- Step 3: Analyze with Context --- 
        root = tk.Tk()
        root.withdraw()
        processing_window = tk.Toplevel(root)
        processing_window.title("Processing")
        processing_window.geometry("300x100")
        root.eval(f'tk::PlaceWindow {str(processing_window)} center')
        tk.Label(processing_window, text="Step 3: Analyzing with Context...").pack(pady=20)
        processing_window.update()

        # Pass the actual search results to the analysis function
        final_analysis = get_gemini_analysis_with_context(api_key, image_path, question_text, search_results_text)
        
        if processing_window:
            processing_window.destroy()
            processing_window = None

        if final_analysis is None:
            messagebox.showerror("API Error", "Failed to get final analysis from Gemini API. Check console.")
            return

        # Show preview 
        preview_response(image_path, final_analysis)

    except Exception as e:
        if processing_window:
            try: processing_window.destroy() 
            except: pass
        messagebox.showerror("Error", f"An unexpected error occurred during step 3: {str(e)}")
        print(f"Error in continue_processing_single_file (Step 3): {str(e)}")

def create_word_document(image_path, gemini_response):
    """Create a Word document with the Gemini response."""
    doc = Document()
    
    # Document title
    title = doc.add_heading('Gemini Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the image
    doc.add_paragraph("Original Question:")
    try:
        max_width = Inches(6)
        doc.add_picture(image_path, width=max_width)
    except Exception as e:
        doc.add_paragraph(f"[Error including image: {str(e)}]")
    
    # Add Gemini's analysis
    doc.add_heading('Gemini Analysis:', 1)
    
    # Extract the content from the Gemini response
    try:
        content = gemini_response["choices"][0]["message"]["content"]
        
        # Split the content by lines and add to document with appropriate formatting
        current_heading = None
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                # This is a main heading
                current_heading = line[2:]
                doc.add_heading(current_heading, 1)
            elif line.startswith('## '):
                # This is a subheading
                current_heading = line[3:]
                doc.add_heading(current_heading, 2)
            else:
                # Regular text - add with appropriate indentation based on context
                p = doc.add_paragraph(line)
                if current_heading and ("Why other answers" in current_heading):
                    p.paragraph_format.left_indent = Inches(0.25)
    except Exception as e:
        doc.add_paragraph(f"Error formatting Gemini response: {str(e)}")
        doc.add_paragraph(str(gemini_response))
    
    # Save the document with the same name as the original file but with docx extension
    base_name = os.path.splitext(image_path)[0]
    output_path = f"{base_name}_analysis.docx"
    
    doc.save(output_path)
    return output_path

def preview_response(image_path, gemini_response):
    """Show a preview of the Gemini response."""
    root = tk.Tk()
    root.title("Gemini Analysis Preview")
    root.geometry("800x600")
    
    # Create a frame for the image
    image_frame = tk.Frame(root)
    image_frame.pack(padx=10, pady=10)
    
    # Load and display the image
    try:
        img = Image.open(image_path)
        # Resize image to fit window if too large
        max_width, max_height = 400, 300
        img.thumbnail((max_width, max_height))
        
        photo = ImageTk.PhotoImage(img)
        image_label = tk.Label(image_frame, image=photo)
        image_label.image = photo  # Keep a reference
        image_label.pack()
    except Exception as e:
        tk.Label(image_frame, text=f"Error loading image: {str(e)}").pack()
    
    # Create a frame for the response
    response_frame = tk.Frame(root)
    response_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    # Add a scrollable text widget for the response
    response_text = tk.Text(response_frame, wrap=tk.WORD, padx=10, pady=10)
    scrollbar = tk.Scrollbar(response_frame, command=response_text.yview)
    response_text.configure(yscrollcommand=scrollbar.set)
    
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    response_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    # Insert the Gemini response
    try:
        content = gemini_response["choices"][0]["message"]["content"]
        response_text.insert(tk.END, content)
    except Exception as e:
        response_text.insert(tk.END, f"Error displaying Gemini response: {str(e)}\n\n{str(gemini_response)}")
    
    # Disable editing
    response_text.configure(state=tk.DISABLED)
    
    # Create a frame for buttons
    button_frame = tk.Frame(root)
    button_frame.pack(pady=10)
    
    # Save button
    save_button = tk.Button(
        button_frame, 
        text="Save to Word Document",
        command=lambda: [
            create_word_document(image_path, gemini_response),
            messagebox.showinfo("Success", f"Analysis saved as {os.path.splitext(image_path)[0]}_analysis.docx"),
            root.destroy()
        ]
    )
    save_button.pack(side=tk.LEFT, padx=10)
    
    # Cancel button
    cancel_button = tk.Button(button_frame, text="Cancel", command=root.destroy)
    cancel_button.pack(side=tk.LEFT, padx=10)
    
    root.mainloop()

def process_folder():
    """Process all image files in a folder."""
    # This is a placeholder for future implementation
    config = load_config()
    
    root = tk.Tk()
    root.withdraw()
    
    initial_dir = config['Settings']['last_folder']
    
    # Ensure the initial directory exists
    if not os.path.exists(initial_dir):
        initial_dir = os.path.expanduser("~")
    
    folder_path = filedialog.askdirectory(
        title="Select folder containing images",
        initialdir=initial_dir
    )
    
    if not folder_path:
        print("No folder selected.")
        return
    
    # Update last folder in config
    config['Settings']['last_folder'] = folder_path
    save_config(config)
    
    # Launch the batch processor
    messagebox.showinfo("Info", "Batch processing needs update for RAG. Please use single file mode.")

def main():
    """Main function."""
    root = tk.Tk()
    root.title("Gemini Image Analyzer")
    root.geometry("400x200")
    
    tk.Label(root, text="Gemini Image Analyzer", font=("Arial", 16)).pack(pady=20)
    
    button_frame = tk.Frame(root)
    button_frame.pack(pady=10)
    
    tk.Button(
        button_frame, 
        text="Process Single Image",
        command=lambda: [root.destroy(), process_single_file()]
    ).pack(pady=10)
    
    tk.Button(
        button_frame,
        text="Process Folder (Use Batch Script)",
        command=lambda: [root.destroy(), process_folder()]
    ).pack(pady=10)
    
    root.mainloop()

if __name__ == "__main__":
    main() 