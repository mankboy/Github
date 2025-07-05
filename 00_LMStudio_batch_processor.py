def preprocess_image(
    self,
    image,
    enhancement_level="Heavy",
    debug=False,
    remove_circles=True,
    debug_save_path=None
):
    """
    Apply preprocessing to improve OCR quality with focus on text clarity.
    - remove_circles: Optionally remove option bubbles/circles.
    - debug_save_path: If provided, saves debug images for inspection.
    """
    import cv2
    import numpy as np
    from PIL import Image

    img_array = np.array(image)
    debug_images = []

    # Convert to grayscale if necessary
    if len(img_array.shape) == 3 and img_array.shape[2] == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    else:
        gray = img_array

    if debug:
        debug_images.append(("Grayscale", Image.fromarray(gray)))

    # Improved dark mode detection and inversion
    mean_val = np.mean(gray)
    if mean_val < 110:  # Lower threshold for very dark backgrounds
        gray = cv2.bitwise_not(gray)
        if debug:
            debug_images.append(("Inverted", Image.fromarray(gray)))

    # Optional: Remove option circles (MCQ bubbles)
    if remove_circles:
        try:
            img_no_circles = gray.copy()
            circles_found = False
            # Try multiple parameter sets for HoughCircles
            for params in [
                (1, 20, 50, 30, 8, 25),
                (1, 20, 30, 20, 5, 30),
                (1.5, 30, 100, 25, 10, 20)
            ]:
                circles = cv2.HoughCircles(
                    gray, cv2.HOUGH_GRADIENT, *params
                )
                if circles is not None:
                    circles = np.uint16(np.around(circles))
                    for x, y, r in circles[0, :]:
                        cv2.circle(img_no_circles, (x, y), r + 2, 255, -1)
                    circles_found = True
                    if debug:
                        debug_images.append(("Circles Removed", Image.fromarray(img_no_circles)))
                    break
            if circles_found:
                gray = img_no_circles
        except Exception as e:
            print(f"Circle removal failed: {e}")

    # Adaptive resizing: scale up small images (max 4x, max 3500px)
    max_dim = max(gray.shape)
    scale = min(4.0, 3500 / max_dim)
    if scale > 1.0:
        gray = cv2.resize(gray, (int(gray.shape[1] * scale), int(gray.shape[0] * scale)), interpolation=cv2.INTER_CUBIC)
        if debug:
            debug_images.append(("Resized", Image.fromarray(gray)))

    # Denoising: strong noise removal but preserve edges
    denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
    if debug:
        debug_images.append(("Denoised", Image.fromarray(denoised)))

    # Ensure image is 2D and uint8 before CLAHE
    if len(denoised.shape) == 3:
        denoised = cv2.cvtColor(denoised, cv2.COLOR_BGR2GRAY)
    if denoised.dtype != np.uint8:
        denoised = denoised.astype(np.uint8)
    # Contrast enhancement: CLAHE with flexible clipLimit
    clahe = cv2.createCLAHE(clipLimit=2.5 if enhancement_level == "Heavy" else 1.5, tileGridSize=(8, 8))
    contrasted = clahe.apply(denoised)
    if debug:
        debug_images.append(("Contrasted", Image.fromarray(contrasted)))

    # Final binarization and morphology
    if enhancement_level == "Light":
        _, result = cv2.threshold(contrasted, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    elif enhancement_level == "Medium":
        binary = cv2.adaptiveThreshold(contrasted, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                       cv2.THRESH_BINARY, 11, 2)
        kernel = np.ones((1, 1), np.uint8)
        result = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
    elif enhancement_level == "Heavy":
        _, binary = cv2.threshold(contrasted, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        # Stronger morphology for tough cases
        kernel = np.ones((2, 2), np.uint8)
        result = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
    else:
        # Default to just contrasted image
        result = contrasted

    if debug:
        debug_images.append(("Final", Image.fromarray(result)))
        # Optionally save debug images for inspection
        if debug_save_path:
            for name, img in debug_images:
                img.save(f"{debug_save_path}_{name}.png")

    return Image.fromarray(result)

import os
import time
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import queue
import configparser
import re
import json
import requests
import base64
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from textblob import TextBlob
import io
import pytesseract
import concurrent.futures
import platform
import subprocess
import shutil
import cv2
import numpy as np

# Version tracking
VERSION = "0.1.1"  # Fix for dark-background screenshots

# Helper function to increment version for developers
def increment_version(version_str=VERSION, level="patch"):
    """
    Increment the version number.
    level: 'major', 'minor', or 'patch'
    """
    major, minor, patch = map(int, version_str.split('.'))
    if level == "major":
        return f"{major + 1}.0.0"
    elif level == "minor":
        return f"{major}.{minor + 1}.0"
    else:  # patch
        return f"{major}.{minor}.{patch + 1}"

def clean_option_text(text):
    # Ensure re and TextBlob are imported at the top of the file
    # Remove leading non-alphanumeric characters and whitespace
    cleaned = re.sub(r"^[^a-zA-Z0-9]+", "", text).strip()
    # Use TextBlob to correct spelling (only if length > 2 and text is not empty)
    if cleaned and len(cleaned.split()) > 0:
        try:
            # Ensure TextBlob is available
            corrected = str(TextBlob(cleaned).correct())
            return corrected
        except Exception: # Broad exception, consider logging or being more specific
            return cleaned # Return original cleaned text on TextBlob error
    return cleaned # Return cleaned text if not processed by TextBlob (e.g. too short or empty)

# Configuration file
CONFIG_FILE = os.path.expanduser("~/.lmstudio_config.ini")

class BatchProcessor:
    def __init__(self, master):
        self.master = master
        master.title(f"LM Studio Batch Processor v{VERSION}")
        
        # Direct Tesseract path for macOS - try multiple approaches at startup
        if platform.system().lower() == "darwin":
            try:
                # First approach: Force pytesseract directly
                pytesseract.pytesseract.tesseract_cmd = "tesseract"
                print("Direct tesseract command set")
                
                # Second approach: Try checking specific path locations
                possible_paths = [
                    "/usr/local/bin/tesseract",
                    "/opt/homebrew/bin/tesseract",
                    "/opt/local/bin/tesseract",
                    # More specific Homebrew locations
                    "/usr/local/Cellar/tesseract/5.5.0_1/bin/tesseract",
                    "/opt/homebrew/Cellar/tesseract/5.5.0_1/bin/tesseract",
                ]
                
                # Try each path
                for path in possible_paths:
                    if os.path.exists(path) and os.access(path, os.X_OK):
                        pytesseract.pytesseract.tesseract_cmd = path
                        print(f"Found tesseract at {path}")
                        break
                
                # Third approach: Try to run tesseract directly to verify it works
                try:
                    version = pytesseract.get_tesseract_version()
                    print(f"Tesseract version: {version}")
                except Exception as e:
                    print(f"Could not get tesseract version: {e}")
                    raise Exception("Tesseract verification failed")
            except Exception as e:
                print(f"Tesseract setup error: {e}")
                master.after(500, lambda: messagebox.showwarning(
                    f"Tesseract Issue - v{VERSION}", 
                    "There was a problem detecting Tesseract. OCR features might not work.\n\n"
                    "Try entering the full path to tesseract in the settings."
                ))
        
        # Load configuration
        self.config = self.load_config()
        
        # Setup UI elements
        self.create_widgets()
        
        # Initialize folder path from config if it exists
        if self.config['Settings'].get('last_folder') and os.path.exists(self.config['Settings']['last_folder']):
            self.folder_var.set(self.config['Settings']['last_folder'])
            self.list_image_files(self.config['Settings']['last_folder'])
        
        # Processing variables
        self.processing_queue = queue.Queue()
        self.results_queue = queue.Queue()
        self.ocr_results = {}  # Store OCR results by file_path
        self.processed_results = [] # Store successful (path, response) tuples
        self.processing_thread = None
        self.stop_processing = False
        self.ocr_thread_pool = None
        
    def load_config(self):
        """Load configuration from file, or create with defaults."""
        config = configparser.ConfigParser()
        
        if os.path.exists(CONFIG_FILE):
            config.read(CONFIG_FILE)
        
        # Ensure required sections and settings exist
        if 'Settings' not in config:
            config['Settings'] = {}
        
        if 'API' not in config:
            config['API'] = {}
            
        if 'last_folder' not in config['Settings']:
            config['Settings']['last_folder'] = os.path.expanduser("~")
            
        if 'endpoint' not in config['API']:
            config['API']['endpoint'] = 'http://localhost:1234/v1/chat/completions'
            
        if 'model' not in config['API']:
            config['API']['model'] = 'gemma-3-12b-instruct'
            
        if 'temperature' not in config['API']:
            config['API']['temperature'] = '0.7'
            
        if 'max_tokens' not in config['API']:
            config['API']['max_tokens'] = '-1'
            
        if 'timeout' not in config['API']:
            config['API']['timeout'] = '300'
            
        if 'max_retries' not in config['API']:
            config['API']['max_retries'] = '2'

        if 'tesseract_path' not in config['Settings']:
            config['Settings']['tesseract_path'] = ''
        
        if 'use_ocr' not in config['Settings']:
            config['Settings']['use_ocr'] = 'False'  # Disabled by default
            
        if 'last_selected_file' not in config['Settings']:
            config['Settings']['last_selected_file'] = ''
        
        self.save_config(config)
        return config
    
    def save_config(self, config=None):
        """Save configuration to file."""
        if config is None:
            config = self.config
            
        # Print for debugging
        print(f"Saving config, last folder: {config['Settings'].get('last_folder', 'Not set')}")
            
        try:
            with open(CONFIG_FILE, 'w') as f:
                config.write(f)
            # Verify config was saved
            if os.path.exists(CONFIG_FILE):
                with open(CONFIG_FILE, 'r') as f:
                    content = f.read()
                    print(f"Config file saved, size: {len(content)} bytes")
        except Exception as e:
            print(f"Error saving config: {str(e)}")
            
    def auto_detect_tesseract(self):
        """Auto-detect Tesseract installation path and update config."""
        detected_path = self.find_tesseract_path()
        if detected_path:
            self.config['Settings']['tesseract_path'] = detected_path
            self.save_config()
            print(f"Tesseract detected at: {detected_path}")
        else:
            print("Tesseract not detected. Path will need to be provided manually or Tesseract installed.")
    
    def find_tesseract_path(self):
        """Find Tesseract installation path based on platform."""
        system = platform.system().lower()
        
        # On macOS, just use the command name and trust PATH
        if system == "darwin":
            return "tesseract"
        
        # For other platforms, continue with detailed detection
        # Check if tesseract is available on PATH
        if self.is_tesseract_on_path():
            return "tesseract" if system != "windows" else "tesseract.exe"
        
        # Platform-specific searches
        if system == "windows":
            # Common Windows installation paths
            common_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ]
            for path in common_paths:
                if os.path.isfile(path):
                    return path
        
        elif system == "darwin":  # macOS
            # Common macOS installation paths (Homebrew, MacPorts)
            common_paths = [
                "/usr/local/bin/tesseract",
                "/opt/local/bin/tesseract",
                "/opt/homebrew/bin/tesseract",
                "/opt/homebrew/Cellar/tesseract/5.5.0_1/bin/tesseract",  # Specific version path
                "/usr/local/Cellar/tesseract/5.5.0_1/bin/tesseract",    # Intel Mac Homebrew
            ]
            
            # Try to find Homebrew's Cellar path dynamically
            try:
                # Check for Homebrew's Cellar directory
                homebrew_locations = ["/opt/homebrew/Cellar", "/usr/local/Cellar"]
                for brew_loc in homebrew_locations:
                    if os.path.exists(brew_loc):
                        tesseract_dirs = os.path.join(brew_loc, "tesseract")
                        if os.path.exists(tesseract_dirs):
                            # Get all version directories (sorted by version)
                            versions = sorted(os.listdir(tesseract_dirs))
                            if versions:
                                # Get the latest version
                                latest = versions[-1]
                                bin_path = os.path.join(tesseract_dirs, latest, "bin", "tesseract")
                                if os.path.exists(bin_path):
                                    common_paths.insert(0, bin_path)  # Add as highest priority
            except Exception as e:
                print(f"Error looking for Homebrew Tesseract: {e}")
            
            # Try to run `brew --prefix` to find Homebrew's prefix
            try:
                result = subprocess.run(
                    ["brew", "--prefix", "tesseract"],
                    capture_output=True,
                    text=True,
                    check=False
                )
                if result.returncode == 0:
                    brew_prefix = result.stdout.strip()
                    if brew_prefix:
                        bin_path = os.path.join(brew_prefix, "bin", "tesseract")
                        if os.path.exists(bin_path):
                            common_paths.insert(0, bin_path)  # Add as highest priority
            except Exception as e:
                print(f"Error running brew --prefix: {e}")
            
            for path in common_paths:
                if os.path.isfile(path):
                    return path
        
        elif system == "linux":
            # Common Linux installation paths
            common_paths = [
                "/usr/bin/tesseract",
                "/usr/local/bin/tesseract"
            ]
            for path in common_paths:
                if os.path.isfile(path):
                    return path
        
        return ""
    
    def is_tesseract_on_path(self):
        """Check if tesseract is available on system PATH."""
        # First try the which command (Unix) or where command (Windows)
        try:
            cmd = "where tesseract" if platform.system().lower() == "windows" else "which tesseract"
            result = subprocess.run(
                cmd, 
                shell=True, 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                check=False
            )
            if result.returncode == 0:
                path = result.stdout.decode('utf-8').strip()
                if path:
                    return True
        except:
            pass
        
        # Fallback to version check
        try:
            cmd = "tesseract --version"
            result = subprocess.run(
                cmd, 
                shell=True, 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                check=False
            )
            return result.returncode == 0
        except:
            return False
    
    def show_tesseract_instructions(self):
        """Show installation instructions for Tesseract."""
        system = platform.system().lower()
        
        if system == "windows":
            instructions = (
                "Tesseract OCR is not installed or couldn't be detected. To install:\n\n"
                "1. Download Tesseract installer from:\n"
                "   https://github.com/UB-Mannheim/tesseract/wiki\n"
                "2. Run the installer and complete the installation\n"
                "3. Either add the installation path to system PATH or\n"
                "   enter the full path (e.g., C:\\Program Files\\Tesseract-OCR\\tesseract.exe)"
            )
        
        elif system == "darwin":  # macOS
            instructions = (
                "Tesseract OCR is not installed or couldn't be detected. To install:\n\n"
                "Using Homebrew:\n"
                "1. Open Terminal\n"
                "2. Run: brew install tesseract\n\n"
                "Or download from: https://github.com/tesseract-ocr/tesseract"
            )
        
        else:  # Linux
            instructions = (
                "Tesseract OCR is not installed or couldn't be detected. To install:\n\n"
                "Ubuntu/Debian:\n"
                "1. Open Terminal\n"
                "2. Run: sudo apt-get install tesseract-ocr\n\n"
                "Fedora/RHEL/CentOS:\n"
                "1. Open Terminal\n"
                "2. Run: sudo dnf install tesseract"
            )
        
        messagebox.showinfo(f"Tesseract Installation - v{VERSION}", instructions)
    
    def create_widgets(self):
        # Top frame for folder selection
        folder_frame = tk.Frame(self.master)
        folder_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(folder_frame, text="Folder:").pack(side=tk.LEFT)
        
        self.folder_var = tk.StringVar()
        folder_entry = tk.Entry(folder_frame, textvariable=self.folder_var, width=50)
        folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        browse_button = tk.Button(folder_frame, text="Browse", command=self.browse_folder)
        browse_button.pack(side=tk.LEFT, padx=5)
        
        # Version and about button
        about_button = tk.Button(folder_frame, text="About", command=self.show_about_info)
        about_button.pack(side=tk.LEFT, padx=5)
        
        # OCR settings frame
        ocr_frame = tk.Frame(self.master)
        ocr_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Tesseract path setting
        tesseract_frame = tk.Frame(ocr_frame)
        tesseract_frame.pack(fill=tk.X, pady=2)
        tk.Label(tesseract_frame, text="Tesseract Path:").pack(side=tk.LEFT)
        self.tesseract_var = tk.StringVar(value=self.config['Settings']['tesseract_path'])
        tesseract_entry = tk.Entry(tesseract_frame, textvariable=self.tesseract_var, width=40)
        tesseract_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # Add help button for Tesseract
        help_button = tk.Button(tesseract_frame, text="?", command=self.show_tesseract_instructions, width=2)
        help_button.pack(side=tk.LEFT, padx=2)
        
        # Test OCR button
        test_button = tk.Button(tesseract_frame, text="Test OCR", command=self.test_ocr)
        test_button.pack(side=tk.LEFT, padx=5)
        
        # Use OCR checkbox
        self.use_ocr_var = tk.BooleanVar(value=self.config['Settings']['use_ocr'].lower() == 'true')
        ocr_cb = tk.Checkbutton(tesseract_frame, text="Use OCR", variable=self.use_ocr_var)
        ocr_cb.pack(side=tk.LEFT, padx=5)
        
        # OCR Advanced settings
        ocr_adv_frame = tk.Frame(ocr_frame)
        ocr_adv_frame.pack(fill=tk.X, pady=2)
        
        # OCR Enhancement Level
        tk.Label(ocr_adv_frame, text="Enhancement:").pack(side=tk.LEFT)
        self.ocr_enhance_var = tk.StringVar(value="Medium")
        enhance_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.ocr_enhance_var, 
                                   values=["None", "Light", "Medium", "Heavy", "Adaptive", "Super"], 
                                   width=10, state="readonly")
        enhance_combo.pack(side=tk.LEFT, padx=5)
        
        # OCR Page Segmentation Mode
        tk.Label(ocr_adv_frame, text="Layout:").pack(side=tk.LEFT)
        self.psm_var = tk.StringVar(value="Auto")
        psm_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.psm_var, 
                                values=["Auto", "Single Block", "Single Column", "Single Line", "Single Word", "Sparse Text"], 
                                width=12, state="readonly")
        psm_combo.pack(side=tk.LEFT, padx=5)
        
        # OCR Engine Mode
        tk.Label(ocr_adv_frame, text="Engine:").pack(side=tk.LEFT)
        self.oem_var = tk.StringVar(value="Default")
        oem_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.oem_var, 
                               values=["Legacy", "Neural", "Both", "Default"], 
                               width=8, state="readonly")
        oem_combo.pack(side=tk.LEFT, padx=5)
        
        # OCR Language
        tk.Label(ocr_adv_frame, text="Language:").pack(side=tk.LEFT)
        self.lang_var = tk.StringVar(value="eng")
        lang_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.lang_var, 
                                values=["eng", "eng+math", "eng+equ"], 
                                width=8, state="readonly")
        lang_combo.pack(side=tk.LEFT, padx=5)
        
        # Show preprocessing checkbox
        self.show_preprocess_var = tk.BooleanVar(value=False)
        preprocess_cb = tk.Checkbutton(ocr_adv_frame, text="Show Preprocessing", variable=self.show_preprocess_var)
        preprocess_cb.pack(side=tk.LEFT, padx=5)
        
        # API settings frame
        api_frame = tk.Frame(self.master)
        api_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Model selection
        model_frame = tk.Frame(api_frame)
        model_frame.pack(fill=tk.X, pady=2)
        tk.Label(model_frame, text="Model:").pack(side=tk.LEFT)
        self.model_var = tk.StringVar(value=self.config['API']['model'])
        model_entry = tk.Entry(model_frame, textvariable=self.model_var, width=30)
        model_entry.pack(side=tk.LEFT, padx=5)
        
        # Temperature setting
        temp_frame = tk.Frame(api_frame)
        temp_frame.pack(fill=tk.X, pady=2)
        tk.Label(temp_frame, text="Temperature:").pack(side=tk.LEFT)
        self.temp_var = tk.StringVar(value=self.config['API']['temperature'])
        temp_entry = tk.Entry(temp_frame, textvariable=self.temp_var, width=5)
        temp_entry.pack(side=tk.LEFT, padx=5)
        
        # Max tokens setting
        tokens_frame = tk.Frame(api_frame)
        tokens_frame.pack(fill=tk.X, pady=2)
        tk.Label(tokens_frame, text="Max Tokens:").pack(side=tk.LEFT)
        self.tokens_var = tk.StringVar(value=self.config['API']['max_tokens'])
        tokens_entry = tk.Entry(tokens_frame, textvariable=self.tokens_var, width=5)
        tokens_entry.pack(side=tk.LEFT, padx=5)
        
        # Timeout setting
        timeout_frame = tk.Frame(api_frame)
        timeout_frame.pack(fill=tk.X, pady=2)
        tk.Label(timeout_frame, text="Timeout (sec):").pack(side=tk.LEFT)
        self.timeout_var = tk.StringVar(value=self.config['API']['timeout'])
        timeout_entry = tk.Entry(timeout_frame, textvariable=self.timeout_var, width=5)
        timeout_entry.pack(side=tk.LEFT, padx=5)
        
        # Max retries setting
        retry_frame = tk.Frame(api_frame)
        retry_frame.pack(fill=tk.X, pady=2)
        tk.Label(retry_frame, text="Max Retries:").pack(side=tk.LEFT)
        self.retry_var = tk.StringVar(value=self.config['API']['max_retries'])
        retry_entry = tk.Entry(retry_frame, textvariable=self.retry_var, width=2)
        retry_entry.pack(side=tk.LEFT, padx=5)
        
        # Vision model checkbox
        vision_frame = tk.Frame(api_frame)
        vision_frame.pack(fill=tk.X, pady=2)
        self.vision_var = tk.BooleanVar(value=False)
        vision_cb = tk.Checkbutton(vision_frame, text="Model has vision capabilities", variable=self.vision_var)
        vision_cb.pack(side=tk.LEFT, padx=5)
        
        # Frame for file list
        files_frame = tk.Frame(self.master)
        files_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tk.Label(files_frame, text="Image files to process:").pack(anchor=tk.W)
        
        # Create a frame with scrollbars for the file list
        file_list_frame = tk.Frame(files_frame)
        file_list_frame.pack(fill=tk.BOTH, expand=True)
        
        # Add scrollbars to the file list
        scrollbar_y = tk.Scrollbar(file_list_frame)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        scrollbar_x = tk.Scrollbar(file_list_frame, orient=tk.HORIZONTAL)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.file_listbox = tk.Listbox(
            file_list_frame,
            selectmode=tk.EXTENDED,
            yscrollcommand=scrollbar_y.set,
            xscrollcommand=scrollbar_x.set
        )
        self.file_listbox.pack(fill=tk.BOTH, expand=True)
        
        scrollbar_y.config(command=self.file_listbox.yview)
        scrollbar_x.config(command=self.file_listbox.xview)
        
        # Control buttons
        control_frame = tk.Frame(self.master)
        control_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.process_button = tk.Button(
            control_frame, 
            text="Process & Create Combined Doc", 
            command=self.start_processing
        )
        self.process_button.pack(side=tk.LEFT, padx=5)
        
        self.stop_button = tk.Button(
            control_frame, 
            text="Stop", 
            command=self.stop_processing_thread,
            state=tk.DISABLED
        )
        self.stop_button.pack(side=tk.LEFT, padx=5)
        
        # Add OCR single-image button
        self.ocr_single_button = tk.Button(
            control_frame,
            text="Run OCR on Selected Image",
            command=self.run_ocr_on_single_image
        )
        self.ocr_single_button.pack(side=tk.LEFT, padx=5)
        
        # Add Quit button
        self.quit_button = tk.Button(
            control_frame,
            text="Quit",
            command=self.quit_application
        )
        self.quit_button.pack(side=tk.RIGHT, padx=5)
        
        # Progress frame
        progress_frame = tk.Frame(self.master)
        progress_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(progress_frame, text="Progress:").pack(anchor=tk.W)
        
        self.progress_var = tk.StringVar(value="Ready")
        tk.Label(progress_frame, textvariable=self.progress_var).pack(anchor=tk.W)
        
        self.progress_bar = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=100, mode='determinate')
        self.progress_bar.pack(fill=tk.X, pady=5)

    def run_ocr_on_single_image(self):
        """Minimal, robust OCR runner for a single image with GUI feedback."""
        from tkinter import filedialog, messagebox
        import traceback
        try:
            file_path = filedialog.askopenfilename(
                title="Select an image for OCR",
                filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.tiff"), ("All files", "*.*")]
            )
            if not file_path:
                return
            self.master.config(cursor="wait")
            self.master.update()
            try:
                text = self.perform_ocr(file_path)
                if not text.strip():
                    message = "No text detected by OCR."
                else:
                    message = text
                messagebox.showinfo("OCR Result", message)
            except Exception as e:
                tb = traceback.format_exc()
                messagebox.showerror("OCR Error", f"Error during OCR:\n{e}\n\n{tb}")
        finally:
            self.master.config(cursor="")
            self.master.update()

        """Minimal, robust OCR runner for a single image with GUI feedback."""
        from tkinter import filedialog, messagebox
        import traceback
        try:
            file_path = filedialog.askopenfilename(
                title="Select an image for OCR",
                filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.tiff"), ("All files", "*.*")]
            )
            if not file_path:
                return
            self.master.config(cursor="wait")
            self.master.update()
            try:
                text = self.perform_ocr(file_path)
                if not text.strip():
                    message = "No text detected by OCR."
                else:
                    message = text
                messagebox.showinfo("OCR Result", message)
            except Exception as e:
                tb = traceback.format_exc()
                messagebox.showerror("OCR Error", f"Error during OCR:\n{e}\n\n{tb}")
        finally:
            self.master.config(cursor="")
            self.master.update()
        # Top frame for folder selection
        folder_frame = tk.Frame(self.master)
        folder_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(folder_frame, text="Folder:").pack(side=tk.LEFT)
        
        self.folder_var = tk.StringVar()
        folder_entry = tk.Entry(folder_frame, textvariable=self.folder_var, width=50)
        folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        browse_button = tk.Button(folder_frame, text="Browse", command=self.browse_folder)
        browse_button.pack(side=tk.LEFT, padx=5)
        
        # Version and about button
        about_button = tk.Button(folder_frame, text="About", command=self.show_about_info)
        about_button.pack(side=tk.LEFT, padx=5)
        
        # OCR settings frame
        ocr_frame = tk.Frame(self.master)
        ocr_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Tesseract path setting
        tesseract_frame = tk.Frame(ocr_frame)
        tesseract_frame.pack(fill=tk.X, pady=2)
        tk.Label(tesseract_frame, text="Tesseract Path:").pack(side=tk.LEFT)
        self.tesseract_var = tk.StringVar(value=self.config['Settings']['tesseract_path'])
        tesseract_entry = tk.Entry(tesseract_frame, textvariable=self.tesseract_var, width=40)
        tesseract_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # Add help button for Tesseract
        help_button = tk.Button(tesseract_frame, text="?", command=self.show_tesseract_instructions, width=2)
        help_button.pack(side=tk.LEFT, padx=2)
        
        # Test OCR button
        test_button = tk.Button(tesseract_frame, text="Test OCR", command=self.test_ocr)
        test_button.pack(side=tk.LEFT, padx=5)
        
        # Use OCR checkbox
        self.use_ocr_var = tk.BooleanVar(value=self.config['Settings']['use_ocr'].lower() == 'true')
        ocr_cb = tk.Checkbutton(tesseract_frame, text="Use OCR", variable=self.use_ocr_var)
        ocr_cb.pack(side=tk.LEFT, padx=5)
        
        # OCR Advanced settings
        ocr_adv_frame = tk.Frame(ocr_frame)
        ocr_adv_frame.pack(fill=tk.X, pady=2)
        
        # OCR Enhancement Level
        tk.Label(ocr_adv_frame, text="Enhancement:").pack(side=tk.LEFT)
        self.ocr_enhance_var = tk.StringVar(value="Medium")
        enhance_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.ocr_enhance_var, 
                                   values=["None", "Light", "Medium", "Heavy", "Adaptive", "Super"], 
                                   width=10, state="readonly")
        enhance_combo.pack(side=tk.LEFT, padx=5)
        
        # OCR Page Segmentation Mode
        tk.Label(ocr_adv_frame, text="Layout:").pack(side=tk.LEFT)
        self.psm_var = tk.StringVar(value="Auto")
        psm_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.psm_var, 
                                values=["Auto", "Single Block", "Single Column", "Single Line", "Single Word", "Sparse Text"], 
                                width=12, state="readonly")
        psm_combo.pack(side=tk.LEFT, padx=5)
        
        # OCR Engine Mode
        tk.Label(ocr_adv_frame, text="Engine:").pack(side=tk.LEFT)
        self.oem_var = tk.StringVar(value="Default")
        oem_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.oem_var, 
                               values=["Legacy", "Neural", "Both", "Default"], 
                               width=8, state="readonly")
        oem_combo.pack(side=tk.LEFT, padx=5)
        
        # OCR Language
        tk.Label(ocr_adv_frame, text="Language:").pack(side=tk.LEFT)
        self.lang_var = tk.StringVar(value="eng")
        lang_combo = ttk.Combobox(ocr_adv_frame, textvariable=self.lang_var, 
                                values=["eng", "eng+math", "eng+equ"], 
                                width=8, state="readonly")
        lang_combo.pack(side=tk.LEFT, padx=5)
        
        # Show preprocessing checkbox
        self.show_preprocess_var = tk.BooleanVar(value=False)
        preprocess_cb = tk.Checkbutton(ocr_adv_frame, text="Show Preprocessing", variable=self.show_preprocess_var)
        preprocess_cb.pack(side=tk.LEFT, padx=5)
        
        # API settings frame
        api_frame = tk.Frame(self.master)
        api_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Model selection
        model_frame = tk.Frame(api_frame)
        model_frame.pack(fill=tk.X, pady=2)
        tk.Label(model_frame, text="Model:").pack(side=tk.LEFT)
        self.model_var = tk.StringVar(value=self.config['API']['model'])
        model_entry = tk.Entry(model_frame, textvariable=self.model_var, width=30)
        model_entry.pack(side=tk.LEFT, padx=5)
        
        # Temperature setting
        temp_frame = tk.Frame(api_frame)
        temp_frame.pack(fill=tk.X, pady=2)
        tk.Label(temp_frame, text="Temperature:").pack(side=tk.LEFT)
        self.temp_var = tk.StringVar(value=self.config['API']['temperature'])
        temp_entry = tk.Entry(temp_frame, textvariable=self.temp_var, width=5)
        temp_entry.pack(side=tk.LEFT, padx=5)
        
        # Max tokens setting
        tokens_frame = tk.Frame(api_frame)
        tokens_frame.pack(fill=tk.X, pady=2)
        tk.Label(tokens_frame, text="Max Tokens:").pack(side=tk.LEFT)
        self.tokens_var = tk.StringVar(value=self.config['API']['max_tokens'])
        tokens_entry = tk.Entry(tokens_frame, textvariable=self.tokens_var, width=5)
        tokens_entry.pack(side=tk.LEFT, padx=5)
        
        # Timeout setting
        timeout_frame = tk.Frame(api_frame)
        timeout_frame.pack(fill=tk.X, pady=2)
        tk.Label(timeout_frame, text="Timeout (sec):").pack(side=tk.LEFT)
        self.timeout_var = tk.StringVar(value=self.config['API']['timeout'])
        timeout_entry = tk.Entry(timeout_frame, textvariable=self.timeout_var, width=5)
        timeout_entry.pack(side=tk.LEFT, padx=5)
        
        # Max retries setting
        retry_frame = tk.Frame(api_frame)
        retry_frame.pack(fill=tk.X, pady=2)
        tk.Label(retry_frame, text="Max Retries:").pack(side=tk.LEFT)
        self.retry_var = tk.StringVar(value=self.config['API']['max_retries'])
        retry_entry = tk.Entry(retry_frame, textvariable=self.retry_var, width=2)
        retry_entry.pack(side=tk.LEFT, padx=5)
        
        # Vision model checkbox
        vision_frame = tk.Frame(api_frame)
        vision_frame.pack(fill=tk.X, pady=2)
        self.vision_var = tk.BooleanVar(value=False)
        vision_cb = tk.Checkbutton(vision_frame, text="Model has vision capabilities", variable=self.vision_var)
        vision_cb.pack(side=tk.LEFT, padx=5)
        
        # Frame for file list
        files_frame = tk.Frame(self.master)
        files_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tk.Label(files_frame, text="Image files to process:").pack(anchor=tk.W)
        
        # Create a frame with scrollbars for the file list
        file_list_frame = tk.Frame(files_frame)
        file_list_frame.pack(fill=tk.BOTH, expand=True)
        
        # Add scrollbars to the file list
        scrollbar_y = tk.Scrollbar(file_list_frame)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        scrollbar_x = tk.Scrollbar(file_list_frame, orient=tk.HORIZONTAL)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.file_listbox = tk.Listbox(
            file_list_frame,
            selectmode=tk.EXTENDED,
            yscrollcommand=scrollbar_y.set,
            xscrollcommand=scrollbar_x.set
        )
        self.file_listbox.pack(fill=tk.BOTH, expand=True)
        
        scrollbar_y.config(command=self.file_listbox.yview)
        scrollbar_x.config(command=self.file_listbox.xview)
        
        # Control buttons
        control_frame = tk.Frame(self.master)
        control_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.process_button = tk.Button(
            control_frame, 
            text="Process & Create Combined Doc", 
            command=self.start_processing
        )
        self.process_button.pack(side=tk.LEFT, padx=5)
        
        self.stop_button = tk.Button(
            control_frame, 
            text="Stop", 
            command=self.stop_processing_thread,
            state=tk.DISABLED
        )
        self.stop_button.pack(side=tk.LEFT, padx=5)
        
        # Add OCR single-image button
        self.ocr_single_button = tk.Button(
            control_frame,
            text="Run OCR on Selected Image",
            command=self.run_ocr_on_single_image
        )
        self.ocr_single_button.pack(side=tk.LEFT, padx=5)

        # Add Quit button
        self.quit_button = tk.Button(
            control_frame,
            text="Quit",
            command=self.quit_application
        )
        self.quit_button.pack(side=tk.RIGHT, padx=5)
        
        # Progress frame
        progress_frame = tk.Frame(self.master)
        progress_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(progress_frame, text="Progress:").pack(anchor=tk.W)
        
        self.progress_var = tk.StringVar(value="Ready")
        tk.Label(progress_frame, textvariable=self.progress_var).pack(anchor=tk.W)
        
        self.progress_bar = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=100, mode='determinate')
        self.progress_bar.pack(fill=tk.X, pady=5)
    
    def browse_folder(self):
        """Open a folder browser dialog."""
        initial_dir = self.config['Settings'].get('last_folder', os.path.expanduser("~"))
        
        # Ensure the initial directory exists
        if not os.path.exists(initial_dir):
            initial_dir = os.path.expanduser("~")
        
        folder_path = filedialog.askdirectory(
            title="Select folder containing image files",
            initialdir=initial_dir
        )
        
        if folder_path:
            # Update last folder in config and save immediately
            self.config['Settings']['last_folder'] = folder_path
            self.save_config()
            
            # Update the entry and list the files
            self.folder_var.set(folder_path)
            self.list_image_files(folder_path)
    
    def extract_timestamp_from_filename(self, filename):
        """Extract timestamp from filename for sorting."""
        # Common timestamp patterns in filenames
        patterns = [
            r'(\d{4}-\d{2}-\d{2}[-_]\d{2}-\d{2}-\d{2})',  # YYYY-MM-DD-HH-MM-SS
            r'(\d{4}\d{2}\d{2}[-_]\d{2}\d{2}\d{2})',      # YYYYMMDD-HHMMSS
            r'(\d{8}[-_]\d{6})',                         # YYYYMMDD-HHMMSS
            r'(\d{4}-\d{2}-\d{2})',                      # YYYY-MM-DD
            r'(\d{8})',                                  # YYYYMMDD
            r'[-_](\d{6})[-_]',                          # _HHMMSS_
            r'[-_](\d{4})[-_]',                          # _MMSS_
            r'[-_](\d{10,14})[-_]',                      # Unix timestamp
        ]
        
        # Try each pattern
        for pattern in patterns:
            match = re.search(pattern, filename)
            if match:
                timestamp_str = match.group(1)
                # Try to convert to a numeric value for comparison
                # Remove non-numeric characters
                numeric_timestamp = re.sub(r'\D', '', timestamp_str)
                try:
                    return int(numeric_timestamp)
                except ValueError:
                    pass
        
        # Fallback - if no timestamp found, return 0
        return 0

    def list_image_files(self, folder_path):
        """List all image files in the selected folder, sorted by timestamp in filename (oldest first)."""
        self.file_listbox.delete(0, tk.END)
        
        if not folder_path or not os.path.exists(folder_path):
            return
        
        image_files = []
        
        # Collect image files with timestamps from filenames
        for file in os.listdir(folder_path):
            if file.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif')):
                # Extract timestamp from filename
                timestamp = self.extract_timestamp_from_filename(file)
                image_files.append((file, timestamp))
        
        # Sort by timestamp, oldest first (ascending order)
        image_files.sort(key=lambda x: x[1])
        
        # Insert sorted files into listbox
        for file, _ in image_files:
            self.file_listbox.insert(tk.END, file)
        
        # Extract just the filenames for later use
        sorted_filenames = [file for file, _ in image_files]
        
        # Select last used file if it exists
        last_file = self.config['Settings'].get('last_selected_file', '')
        if last_file and last_file in sorted_filenames:
            index = sorted_filenames.index(last_file)
            self.file_listbox.selection_set(index)
            self.file_listbox.see(index)
    
    def start_processing(self):
        """Start processing the image files in a separate thread."""
        folder_path = self.folder_var.get()
        
        if not folder_path or not os.path.exists(folder_path):
            messagebox.showerror("Error", "Please select a valid folder.")
            return
        
        # Update settings in config
        self.config['API']['model'] = self.model_var.get()
        self.config['API']['temperature'] = self.temp_var.get()
        self.config['API']['max_tokens'] = self.tokens_var.get()
        self.config['API']['timeout'] = self.timeout_var.get()
        self.config['API']['max_retries'] = self.retry_var.get()
        self.config['Settings']['tesseract_path'] = self.tesseract_var.get()
        self.config['Settings']['last_folder'] = folder_path
        self.config['Settings']['use_ocr'] = str(self.use_ocr_var.get())
        self.save_config()
        
        # Get all selected files, or all files if none selected
        selected_indices = self.file_listbox.curselection()
        if selected_indices:
            selected_files = [self.file_listbox.get(i) for i in selected_indices]
        else:
            selected_files = [self.file_listbox.get(i) for i in range(self.file_listbox.size())]
        
        if not selected_files:
            messagebox.showerror("Error", "No image files found in the selected folder.")
            return
        
        # Check OCR settings if enabled
        if self.use_ocr_var.get():
            tesseract_path = self.tesseract_var.get()
            if tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Reset results
        self.processed_results = []
        self.ocr_results = {}
        
        # Clear queues
        while not self.processing_queue.empty():
            self.processing_queue.get()
        
        while not self.results_queue.empty():
            self.results_queue.get()
        
        # Create a thread pool for OCR
        max_workers = min(os.cpu_count() or 2, 4)  # Use at most 4 workers
        self.ocr_thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=max_workers)
        
        # Add files to processing queue
        for file in selected_files:
            file_path = os.path.join(folder_path, file)
            self.processing_queue.put(file_path)
            
            # If OCR is enabled, submit OCR task to thread pool
            if self.use_ocr_var.get():
                self.ocr_thread_pool.submit(self.ocr_worker, file_path)
        
        # Start processing thread
        self.stop_processing = False
        self.processing_thread = threading.Thread(
            target=self.process_files, 
            args=(len(selected_files),)
        )
        self.processing_thread.daemon = True
        self.processing_thread.start()
        
        # Update UI
        self.process_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)
        self.progress_var.set("Processing...")
        self.progress_bar["value"] = 0
        self.progress_bar["maximum"] = len(selected_files)
        
        # Start checking for results
        self.master.after(100, self.check_progress)
    
    def ocr_worker(self, file_path):
        """Worker function for OCR processing in thread pool."""
        try:
            file_name = os.path.basename(file_path)
            self.results_queue.put(("status", f"Performing OCR on {file_name} with API..."))
            
            ocr_text = self.perform_ocr(file_path)
            
            # Store the OCR result
            self.ocr_results[file_path] = ocr_text
            self.results_queue.put(("ocr_result", (file_path, ocr_text)))
            
            self.results_queue.put(("status", f"OCR completed for {file_name}"))
            return ocr_text
        except Exception as e:
            self.results_queue.put(("error", f"OCR Error for {os.path.basename(file_path)}: {str(e)}"))
            return f"OCR Error: {str(e)}"
    
    def perform_ocr(self, image_path):
        """Extract text from an image using OCR with multiple approaches."""
        try:
            print(f"OCR processing: {image_path}")
            
            # Open the image with PIL
            img = Image.open(image_path)
            print(f"Image opened: {img.format}, {img.size}")
            
            # Get OCR settings
            enhancement = self.ocr_enhance_var.get()
            psm_name = self.psm_var.get()
            engine_mode = self.oem_var.get()
            language = self.lang_var.get()
            show_preprocessing = self.show_preprocess_var.get()
            
            # Convert PSM name to Tesseract PSM code
            psm_codes = {
                "Auto": 3,
                "Single Block": 6,
                "Single Column": 4,
                "Single Line": 7,
                "Single Word": 8,
                "Sparse Text": 11
            }
            psm = psm_codes.get(psm_name, 3)  # Default to Auto (3) if not found
            
            # Convert OEM name to Tesseract OEM code
            oem_codes = {
                "Legacy": 0,
                "Neural": 1,
                "Both": 2,
                "Default": 3
            }
            oem = oem_codes.get(engine_mode, 3)  # Default to Default (3) if not found
            
            # Add padding to ensure we don't lose content at edges
            padded_img = ImageOps.expand(img, border=20, fill='white')
            
            # Show original image in debug window if enabled
            if show_preprocessing:
                debug_images = [("Original", img)]
            
            # Try different preprocessing approaches
            preprocessed_images = []
            
            # Special handling for dark screenshots (like macOS dark mode)
            # Check if the image is likely a dark mode screenshot
            is_dark_screenshot = self.is_dark_screenshot(img)
            
            # If using Super enhancement or detected dark screenshot, include special handling
            if enhancement == "Super" or is_dark_screenshot:
                # For Super enhancement or dark screenshots, include special processing for dark backgrounds
                special_processed = self.preprocess_dark_screenshot(padded_img, show_preprocessing)
                preprocessed_images.append(("Dark Mode", special_processed))
                
                # Also try inverting the image first and then processing
                inverted_img = ImageOps.invert(padded_img.convert('RGB')).convert('L')
                inverted_processed = self.preprocess_image(inverted_img, "Heavy", show_preprocessing)
                preprocessed_images.append(("Inverted", inverted_processed))
                
                # Try direct OCR on the original image
                direct_img = padded_img.convert('L')  # Convert to grayscale
                preprocessed_images.append(("Direct", direct_img))
            
            # Preprocess the image with selected enhancement level
            if enhancement == "Super":
                # For Super enhancement, try multiple processing methods
                for sub_enhancement in ["Medium", "Heavy", "Adaptive"]:
                    processed = self.preprocess_image(padded_img, sub_enhancement, show_preprocessing)
                    preprocessed_images.append((sub_enhancement, processed))
            else:
                # Use the selected enhancement level
                processed_img = self.preprocess_image(padded_img, enhancement, show_preprocessing)
                preprocessed_images.append((enhancement, processed_img))
            
            # Show preprocessing result if enabled
            if show_preprocessing and hasattr(self, 'master'):
                if len(preprocessed_images) == 1:
                    self.show_preprocessed_image(img, preprocessed_images[0][1], debug_images)
                else:
                    # Show multiple preprocessed images for comparison
                    self.show_multiple_preprocessed_images(img, preprocessed_images, debug_images)
            
            best_text = ""
            best_confidence = 0
            best_settings = ""
            
            # Try all preprocessed images with specified PSM
            for preproc_name, processed_img in preprocessed_images:
                try:
                    # Configure OCR parameters
                    config = f'--oem {oem} --psm {psm} -l {language} -c preserve_interword_spaces=1 -c tessedit_do_invert=0'
                    
                    # For dark mode, try with and without inversion
                    if preproc_name in ["Dark Mode", "Inverted"]:
                        # First try with invert disabled
                        text1 = pytesseract.image_to_string(processed_img, config=config)
                        
                        # Then try with invert enabled
                        invert_config = config.replace("tessedit_do_invert=0", "tessedit_do_invert=1")
                        text2 = pytesseract.image_to_string(processed_img, config=invert_config)
                        
                        # Use the longer text
                        text = text1 if len(text1) > len(text2) else text2
                        print(f"Dark mode OCR: text1 len {len(text1)}, text2 len {len(text2)}")
                        
                        # Calculate average confidence for the better text
                        try:
                            if len(text1) > len(text2):
                                data = pytesseract.image_to_data(processed_img, config=config, output_type=pytesseract.Output.DICT)
                            else:
                                data = pytesseract.image_to_data(processed_img, config=invert_config, output_type=pytesseract.Output.DICT)
                            
                            confidences = [conf for conf in data['conf'] if conf != -1]
                            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                        except Exception as e:
                            print(f"Error getting confidence: {e}")
                            avg_confidence = 50 if text else 0  # Default confidence if we got text
                    else:
                        # Run OCR with confidence for regular processing
                        try:
                            data = pytesseract.image_to_data(processed_img, config=config, output_type=pytesseract.Output.DICT)
                            
                            # Calculate average confidence
                            confidences = [conf for conf in data['conf'] if conf != -1]
                            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                            
                            # Extract the text
                            text = pytesseract.image_to_string(processed_img, config=config)
                        except Exception as e:
                            print(f"Error with standard OCR: {e}")
                            # Fallback to basic image_to_string if image_to_data fails
                            text = pytesseract.image_to_string(processed_img, config=config)
                            avg_confidence = 50 if text else 0  # Default confidence if we got text
                    
                    print(f"OCR with {preproc_name}, PSM {psm}, confidence: {avg_confidence:.2f}, text length: {len(text)}")
                    
                    # Keep the result if it's better than what we have
                    if text and (not best_text or avg_confidence > best_confidence):
                        best_text = text
                        best_confidence = avg_confidence
                        best_settings = f"{preproc_name}, PSM {psm_name}, Confidence: {avg_confidence:.2f}%"
                    
                except Exception as e:
                    print(f"OCR with {preproc_name}, PSM {psm} failed: {e}")
                    continue
            
            # If Super enhancement, also try alternative PSM modes
            if enhancement == "Super" or (not best_text and is_dark_screenshot):
                # Try alternative PSM modes with the best preprocessing
                best_preproc = preprocessed_images[0][1]  # Default to first preprocessing
                
                # Try different PSM modes to get best results
                alternative_psm_modes = [3, 4, 6, 11, 1, 7]  # Auto, Single Column, Single Block, Sparse Text, OSD, Single Line
                
                for alt_psm in alternative_psm_modes:
                    if alt_psm == psm:
                        continue  # Skip the already tried PSM
                    
                    try:
                        # Configure OCR parameters
                        config = f'--oem {oem} --psm {alt_psm} -l {language} -c preserve_interword_spaces=1 -c tessedit_do_invert=0'
                        
                        # Add additional configeration for improved accuracy
                        additional_config = ' -c tessedit_char_whitelist="ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,?!()[]{}:;\'"_-+=/\\ " -c textord_heavy_nr=1 -c textord_min_linesize=2'
                        config += additional_config
                        
                        # Try with and without inversion for dark screenshots
                        text1 = pytesseract.image_to_string(best_preproc, config=config)
                        
                        if is_dark_screenshot:
                            # Try with inversion enabled
                            invert_config = config.replace("tessedit_do_invert=0", "tessedit_do_invert=1")
                            text2 = pytesseract.image_to_string(best_preproc, config=invert_config)
                            
                            # Use the longer text
                            text = text1 if len(text1) > len(text2) else text2
                            current_config = config if len(text1) > len(text2) else invert_config
                        else:
                            text = text1
                            current_config = config
                        
                        # Calculate average confidence
                        try:
                            data = pytesseract.image_to_data(best_preproc, config=current_config, output_type=pytesseract.Output.DICT)
                            confidences = [conf for conf in data['conf'] if conf != -1]
                            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                        except:
                            avg_confidence = 50 if text else 0  # Default confidence if we got text
                        
                        print(f"OCR with Alternative PSM {alt_psm}, confidence: {avg_confidence:.2f}, text length: {len(text)}")
                        
                        # Keep the result if it's better than what we have
                        if text and (not best_text or avg_confidence > best_confidence):
                            best_text = text
                            best_confidence = avg_confidence
                            best_settings = f"Alt PSM {alt_psm}, Confidence: {avg_confidence:.2f}%"
                    
                    except Exception as e:
                        print(f"OCR with Alternative PSM {alt_psm} failed: {e}")
                        continue
            
            # Last resort - if still no text, try to extract directly with minimal preprocessing
            if not best_text:
                try:
                    print("Attempting direct screenshot OCR extraction...")
                    # Direct extraction from original with minimal processing
                    text = self.extract_text_from_screenshot(img)
                    if text:
                        best_text = text
                        best_settings = "Direct extraction"
                except Exception as e:
                    print(f"Direct extraction failed: {e}")
            
            # If all approaches failed, provide error message with troubleshooting info
            if not best_text:
                return ("No text detected in the image.\n\n"
                        "Troubleshooting tips:\n"
                        "1. Try the 'Super' enhancement level\n"
                        "2. Check if the image contains actual text and not just graphics\n"
                        "3. Try a different image format (PNG, JPEG)\n"
                        "4. Ensure Tesseract is properly installed and configured")
            
            # Format the result
            formatted_text = self.format_question_options(best_text)
            return f"[Settings: {best_settings}]\n\n{formatted_text}"
            
        except Exception as e:
            print(f"OCR base error: {e}")
            return f"OCR Error: {str(e)}"
    
    def is_dark_screenshot(self, image):
        """Detect if the image is likely a dark mode screenshot."""
        # Convert to grayscale for analysis
        img_array = np.array(image.convert('L'))
        
        # Calculate average brightness
        avg_brightness = np.mean(img_array)
        
        # Check pixel distribution
        dark_pixels = np.sum(img_array < 50)  # Count very dark pixels
        light_pixels = np.sum(img_array > 200)  # Count very light pixels
        
        total_pixels = img_array.size
        dark_ratio = dark_pixels / total_pixels
        light_text_ratio = light_pixels / total_pixels
        
        print(f"Image analysis: brightness {avg_brightness:.1f}, dark ratio {dark_ratio:.3f}, light ratio {light_text_ratio:.3f}")
        
        # If image has low average brightness and a significant amount of dark pixels
        # but also some light pixels (text), it's likely a dark mode screenshot
        is_dark = (avg_brightness < 128 and dark_ratio > 0.5 and light_text_ratio > 0.01)
        
        if is_dark:
            print("Detected dark mode screenshot")
        
        return is_dark
    
    def preprocess_dark_screenshot(self, image, debug=False):
        """Special preprocessing for dark mode screenshots."""
        # Convert to numpy array and RGB to ensure proper processing
        img_array = np.array(image.convert('RGB'))
        
        # Debug images if requested
        debug_images = []
        
        # Convert to grayscale
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        if debug:
            debug_images.append(("Grayscale", Image.fromarray(gray)))
        
        # Invert the image (dark background becomes light, light text becomes dark)
        inverted = cv2.bitwise_not(gray)
        if debug:
            debug_images.append(("Inverted", Image.fromarray(inverted)))
        
        # Apply gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(inverted, (3, 3), 0)
        if debug:
            debug_images.append(("Blurred", Image.fromarray(blurred)))
        
        # Apply adaptive thresholding to enhance text
        adaptive_thresh = cv2.adaptiveThreshold(
            blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )
        if debug:
            debug_images.append(("Adaptive Threshold", Image.fromarray(adaptive_thresh)))
        
        # Remove small noise
        kernel = np.ones((2, 2), np.uint8)
        cleaned = cv2.morphologyEx(adaptive_thresh, cv2.MORPH_CLOSE, kernel)
        if debug:
            debug_images.append(("Cleaned", Image.fromarray(cleaned)))
        
        # Add border for better OCR handling
        bordered = cv2.copyMakeBorder(cleaned, 20, 20, 20, 20, cv2.BORDER_CONSTANT, value=255)
        if debug:
            debug_images.append(("Bordered", Image.fromarray(bordered)))
        
        return Image.fromarray(bordered)
    
    def extract_text_from_screenshot(self, image):
        """Direct extraction specialized for screenshots."""
        # Convert image to RGB
        rgb_img = image.convert('RGB')
        
        # Create binary mask for light text on dark background
        img_array = np.array(rgb_img)
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        
        # Try multiple threshold approaches
        _, binary1 = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY)  # High threshold for light text
        _, binary2 = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)  # Medium threshold
        
        # Try with both directly and with inversion
        text1 = pytesseract.image_to_string(Image.fromarray(binary1))
        text2 = pytesseract.image_to_string(Image.fromarray(binary2))
        text3 = pytesseract.image_to_string(Image.fromarray(cv2.bitwise_not(binary1)))
        text4 = pytesseract.image_to_string(Image.fromarray(cv2.bitwise_not(binary2)))
        
        # Use the longest result
        texts = [text1, text2, text3, text4]
        text_lengths = [len(t) for t in texts]
        
        if max(text_lengths) > 0:
            best_idx = text_lengths.index(max(text_lengths))
            return texts[best_idx]
        
        return ""
    
    def preprocess_image(self, image, enhancement_level="Heavy", debug=False):
        """Apply preprocessing to improve OCR quality with focus on text clarity."""
        # Convert PIL Image to OpenCV format
        img_array = np.array(image)
        
        # Debug images if requested
        debug_images = []
        
        # Convert to grayscale if it's not already
        if len(img_array.shape) == 3 and img_array.shape[2] == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        if debug:
            debug_images.append(("Grayscale", Image.fromarray(gray)))
            
        # Simple check for dark background
        mean_val = np.mean(gray)
        if mean_val < 128:  # Dark background
            gray = cv2.bitwise_not(gray)  # Invert for better processing
            if debug:
                debug_images.append(("Inverted", Image.fromarray(gray)))
            
        # Apply more aggressive preprocessing for all levels except "None"
        if enhancement_level == "None":
            # Even with "None", ensure we have proper image format for OCR
            return Image.fromarray(gray)
        
        # Try to remove circular elements (option circles) from the image
        # These can be detected using HoughCircles or by removing smaller circular contours
        img_without_circles = None
        try:
            # Make a copy for circle removal
            img_without_circles = gray.copy()
            
            # Detect and remove small circles (like multiple choice indicators)
            # Try more parameter combinations for better detection
            circle_params = [
                # (dp, minDist, param1, param2, minRadius, maxRadius)
                (1, 20, 50, 30, 8, 25),  # Original parameters
                (1, 20, 30, 20, 5, 30),  # More sensitive
                (1.5, 30, 100, 25, 10, 20)  # More specific
            ]
            
            circles_found = False
            for dp, minDist, param1, param2, minRadius, maxRadius in circle_params:
                circles = cv2.HoughCircles(
                    gray, 
                    cv2.HOUGH_GRADIENT, 
                    dp=dp, 
                    minDist=minDist, 
                    param1=param1, 
                    param2=param2, 
                    minRadius=minRadius, 
                    maxRadius=maxRadius
                )
                
                if circles is not None:
                    circles = np.uint16(np.around(circles))
                    for circle in circles[0, :]:
                        # Get circle coordinates and radius
                        x, y, r = circle
                        
                        # Draw a white filled circle to mask out the detected circle
                        cv2.circle(img_without_circles, (x, y), r, 255, -1)
                        
                        # Also remove slightly outside the circle (for dark borders)
                        cv2.circle(img_without_circles, (x, y), r+2, 255, 2)
                    
                    circles_found = True
                    if debug:
                        debug_images.append(("Circles Removed", Image.fromarray(img_without_circles)))
                    break
            
            # If circles found, use the version without circles 
            if circles_found:
                gray = img_without_circles
        except Exception as e:
            print(f"Circle removal failed: {e}")
            # Continue with original grayscale image if circle removal fails
            pass
        
        # Resize for better processing (higher resolution)
        # Maintain higher resolution for better detection
        scale_percent = min(2.0, 3000 / max(gray.shape[0], gray.shape[1]))
        if scale_percent != 1.0:
            width = int(gray.shape[1] * scale_percent)
            height = int(gray.shape[0] * scale_percent)
            gray = cv2.resize(gray, (width, height), interpolation=cv2.INTER_CUBIC)
            if debug:
                debug_images.append(("Resized", Image.fromarray(gray)))
        
        # Strong noise removal but preserve edges
        denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
        if debug:
            debug_images.append(("Denoised", Image.fromarray(denoised)))
        
        # Increase contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        contrasted = clahe.apply(denoised)
        if debug:
            debug_images.append(("Contrasted", Image.fromarray(contrasted)))
        
        # Apply different final processing based on enhancement level
        if enhancement_level == "Light":
            # Simple threshold
            _, binary = cv2.threshold(contrasted, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            result = binary
            
        elif enhancement_level == "Medium":
            # Adaptive threshold
            binary = cv2.adaptiveThreshold(contrasted, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                         cv2.THRESH_BINARY, 11, 2)
            
            # Light morphology
            kernel = np.ones((1, 1), np.uint8)
            result = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
            
        elif enhancement_level == "Heavy":
            # Apply Gaussian blur to reduce noise before thresholding
            blurred = cv2.GaussianBlur(contrasted, (3, 3), 0)
            
            # Two-pass thresholding for better results
            _, binary = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # More aggressive morphological operations
            kernel = np.ones((2, 2), np.uint8)
            opened = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel)
            result = cv2.morphologyEx(opened, cv2.MORPH_CLOSE, kernel)
            
        elif enhancement_level == "Adaptive":
            # Use both adaptive and Otsu thresholding and combine results
            adaptive_thresh = cv2.adaptiveThreshold(contrasted, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                                 cv2.THRESH_BINARY, 15, 8)
            
            ret, otsu_thresh = cv2.threshold(contrasted, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Combine the two approaches (bitwise OR)
            combined = cv2.bitwise_or(adaptive_thresh, otsu_thresh)
            
            # Clean up with morphology
            kernel = np.ones((2, 2), np.uint8)
            result = cv2.morphologyEx(combined, cv2.MORPH_CLOSE, kernel)
            
        else:
            # Default case
            _, result = cv2.threshold(contrasted, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        if debug:
            debug_images.append((f"Final ({enhancement_level})", Image.fromarray(result)))
        
        # Add final border to ensure text isn't cut off at edges during OCR
        result = cv2.copyMakeBorder(result, 10, 10, 10, 10, cv2.BORDER_CONSTANT, value=255)
        
        # Convert back to PIL Image
        return Image.fromarray(result)
    
    def show_preprocessed_image(self, original_img, processed_img, debug_images=None):
        """Display the original and preprocessed images side by side."""
        preview_window = tk.Toplevel(self.master)
        preview_window.title(f"OCR Image Preprocessing - v{VERSION}")
        
        # Make the window larger if debug images are included
        if debug_images and len(debug_images) > 0:
            preview_window.geometry("1200x800")
        else:
            preview_window.geometry("800x600")
        
        # Create main frame with scrollbar
        main_frame = tk.Frame(preview_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Add a canvas with scrollbar for many images
        canvas = tk.Canvas(main_frame)
        scrollbar = tk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(
                scrollregion=canvas.bbox("all")
            )
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Create frames for original and final images at the top
        top_row = tk.Frame(scrollable_frame)
        top_row.pack(fill=tk.X, expand=True, pady=5)
        
        frame_orig = tk.Frame(top_row)
        frame_orig.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        frame_proc = tk.Frame(top_row)
        frame_proc.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Add labels
        tk.Label(frame_orig, text="Original Image").pack()
        tk.Label(frame_proc, text="Preprocessed Image").pack()
        
        # Resize images to fit in window
        max_display_size = 350
        
        # Resize original
        width, height = original_img.size
        scale = min(max_display_size / width, max_display_size / height)
        new_size = (int(width * scale), int(height * scale))
        display_original = original_img.resize(new_size)
        
        # Resize processed
        width, height = processed_img.size
        scale = min(max_display_size / width, max_display_size / height)
        new_size = (int(width * scale), int(height * scale))
        display_processed = processed_img.resize(new_size)
        
        # Convert to PhotoImage objects
        from PIL import ImageTk
        photo_orig = ImageTk.PhotoImage(display_original)
        photo_proc = ImageTk.PhotoImage(display_processed)
        
        # Keep references to prevent garbage collection
        preview_window.photo_orig = photo_orig
        preview_window.photo_proc = photo_proc
        
        # Display images
        label_orig = tk.Label(frame_orig, image=photo_orig)
        label_orig.pack(padx=5, pady=5)
        
        label_proc = tk.Label(frame_proc, image=photo_proc)
        label_proc.pack(padx=5, pady=5)
        
        # If we have debug images, display them too
        if debug_images and len(debug_images) > 0:
            # Create a separator
            separator = ttk.Separator(scrollable_frame, orient='horizontal')
            separator.pack(fill='x', pady=10)
            
            tk.Label(scrollable_frame, text="Processing Steps", font=('TkDefaultFont', 12, 'bold')).pack(pady=5)
            
            # Create a frame for debug images
            debug_frame = tk.Frame(scrollable_frame)
            debug_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
            
            # Keep track of all photo references
            preview_window.debug_photos = []
            
            # Create a grid of debug images
            max_cols = 3
            row, col = 0, 0
            
            for name, img in debug_images:
                # Create frame for this debug image
                img_frame = tk.Frame(debug_frame)
                img_frame.grid(row=row, column=col, padx=5, pady=5, sticky="nsew")
                
                # Add label
                tk.Label(img_frame, text=name).pack()
                
                # Resize image
                width, height = img.size
                scale = min(max_display_size / width, max_display_size / height)
                new_size = (int(width * scale), int(height * scale))
                display_img = img.resize(new_size)
                
                # Convert to PhotoImage
                photo = ImageTk.PhotoImage(display_img)
                preview_window.debug_photos.append(photo)
                
                # Display image
                label = tk.Label(img_frame, image=photo)
                label.pack(padx=5, pady=5)
                
                # Update position
                col += 1
                if col >= max_cols:
                    col = 0
                    row += 1
        
        # Add close button
        tk.Button(preview_window, text="Close", command=preview_window.destroy).pack(pady=10)
    
    def show_multiple_preprocessed_images(self, original_img, preprocessed_images, debug_images=None):
        """Display multiple preprocessed images for comparison."""
        preview_window = tk.Toplevel(self.master)
        preview_window.title(f"OCR Image Preprocessing Comparison - v{VERSION}")
        preview_window.geometry("1200x800")
        
        # Create notebook for different tabs
        notebook = ttk.Notebook(preview_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create a tab for original vs each preprocessing
        for preproc_name, processed_img in preprocessed_images:
            # Create a tab
            tab = ttk.Frame(notebook)
            notebook.add(tab, text=preproc_name)
            
            # Create frames for original and processed
            frame_orig = tk.Frame(tab)
            frame_orig.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            frame_proc = tk.Frame(tab)
            frame_proc.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # Add labels
            tk.Label(frame_orig, text="Original Image").pack()
            tk.Label(frame_proc, text=f"Processed ({preproc_name})").pack()
            
            # Resize images to fit in window
            max_display_size = 350
            
            # Resize original
            width, height = original_img.size
            scale = min(max_display_size / width, max_display_size / height)
            new_size = (int(width * scale), int(height * scale))
            display_original = original_img.resize(new_size)
            
            # Resize processed
            width, height = processed_img.size
            scale = min(max_display_size / width, max_display_size / height)
            new_size = (int(width * scale), int(height * scale))
            display_processed = processed_img.resize(new_size)
            
            # Convert to PhotoImage objects
            from PIL import ImageTk
            photo_orig = ImageTk.PhotoImage(display_original)
            photo_proc = ImageTk.PhotoImage(display_processed)
            
            # Store references
            setattr(preview_window, f"photo_orig_{preproc_name}", photo_orig)
            setattr(preview_window, f"photo_proc_{preproc_name}", photo_proc)
            
            # Display images
            label_orig = tk.Label(frame_orig, image=photo_orig)
            label_orig.pack(padx=5, pady=5)
            
            label_proc = tk.Label(frame_proc, image=photo_proc)
            label_proc.pack(padx=5, pady=5)
        
        # If we have debug images, add a tab for the processing steps
        if debug_images and len(debug_images) > 0:
            # Create a tab for debug images
            debug_tab = ttk.Frame(notebook)
            notebook.add(debug_tab, text="Processing Steps")
            
            # Create a canvas with scrollbar
            canvas = tk.Canvas(debug_tab)
            scrollbar = tk.Scrollbar(debug_tab, orient="vertical", command=canvas.yview)
            scrollable_frame = tk.Frame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(
                    scrollregion=canvas.bbox("all")
                )
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # Keep track of all photo references
            preview_window.debug_photos = []
            
            # Create a grid of debug images
            max_cols = 3
            row, col = 0, 0
            
            for name, img in debug_images:
                # Create frame for this debug image
                img_frame = tk.Frame(scrollable_frame)
                img_frame.grid(row=row, column=col, padx=5, pady=5, sticky="nsew")
                
                # Add label
                tk.Label(img_frame, text=name).pack()
                
                # Resize image
                max_display_size = 250
                width, height = img.size
                scale = min(max_display_size / width, max_display_size / height)
                new_size = (int(width * scale), int(height * scale))
                display_img = img.resize(new_size)
                
                # Convert to PhotoImage
                photo = ImageTk.PhotoImage(display_img)
                preview_window.debug_photos.append(photo)
                
                # Display image
                label = tk.Label(img_frame, image=photo)
                label.pack(padx=5, pady=5)
                
                # Update position
                col += 1
                if col >= max_cols:
                    col = 0
                    row += 1
        
        # Add close button
        tk.Button(preview_window, text="Close", command=preview_window.destroy).pack(pady=10)
    
    


    def format_question_options(self, clean_text):
        replacements = {
            "0ption": "Option",
            "0ptions": "Options",
            "Ouestion": "Question",
            "Qption": "Option",
            "0:": "Option:",
            "0.": "Option:",
            "|": "I",  # Common OCR mistake
            "l.": "1.",  # lowercase L vs 1
            "l)": "1)",
            "O.": "0.",  # Letter O vs zero
            "O)": "0)",
            "(A)": "A)",
            "(B)": "B)",
            "(C)": "C)",
            "(D)": "D)",
        }
        for error, correction in replacements.items():
            clean_text = clean_text.replace(error, correction)
        # Remove selection circles and bullet points from the text
        clean_text = re.sub(r'[○●•◉⦿⚪⚫]|[\(\[]?\s*[Oo]\s*[\)\]]?\s*', '', clean_text)
        # First, try to identify a question directly
        question_match = re.search(r'([^.!?]+\?)', clean_text)
        if question_match:
            question_text = question_match.group(1).strip()
            remaining_text = clean_text.replace(question_text, "", 1)
        else:
            lines = clean_text.split('\n')
            question_text = ""
            remaining_text = clean_text
            option_start_index = -1
            for i, line in enumerate(lines):
                if (re.search(r'^\s*A[\.\:\)\s]', line) or
                    re.search(r'^\s*[Oo]?ption\s*A[\.\:\s]', line, re.IGNORECASE) or
                    re.search(r'^\s*A$', line)):
                    option_start_index = i
                    break
            if option_start_index > 0:
                question_text = " ".join(lines[:option_start_index]).strip()
                remaining_text = "\n".join(lines[option_start_index:])
        if question_text and not question_text.endswith('?'):
            question_text = question_text + '?'
        question_text = clean_option_text(question_text)
        options = {'A': '', 'B': '', 'C': '', 'D': ''}
        option_lines = []
        bullet_regex = re.compile(r'^[\s\u2022\u2023\u25E6\u2043\u2219\*\-]+(.+)$')
        lines = [l.strip() for l in remaining_text.split('\n') if l.strip()]
        used_indices = set()
        for option in options.keys():
            patterns = [
                fr'(?:^|\n)\s*{option}[\.\:\)\s]\s*(.*?)(?=(?:\n\s*[B-D][\.\:\)\s])|$)',
                fr'(?:^|\n)\s*[Oo]ption\s+{option}[\.\:\s]\s*(.*?)(?=(?:\n\s*[Oo]ption\s+[B-D])|$)',
                fr'(?:^|\n)\s*[Oo]ptions:\s*{option}[\.\:\s]\s*(.*?)(?=(?:\n\s*[Oo]ptions:\s*[B-D])|$)',
            ]
            for i, line in enumerate(lines):
                for pattern in patterns:
                    match = re.match(pattern, line, re.IGNORECASE | re.DOTALL)
                    if match:
                        options[option] = match.group(1).strip()
                        used_indices.add(i)
                        break
                if options[option]:
                    break
        bullet_lines = [i for i, l in enumerate(lines) if bullet_regex.match(l) and i not in used_indices]
        for idx, option in zip(bullet_lines, options.keys()):
            bullet_match = bullet_regex.match(lines[idx])
            if bullet_match:
                options[option] = bullet_match.group(1).strip()
        empty_options = [k for k, v in options.items() if not v]
        used_lines = set(used_indices) | set(bullet_lines)
        candidate_lines = [l for i, l in enumerate(lines) if i not in used_lines]
        for option, line in zip(empty_options, candidate_lines):
            options[option] = line.strip()
        for option in options:
            options[option] = clean_option_text(options[option])
        result = f"Question: {question_text.strip()}\n\n"
        for option, text in options.items():
            result += f"Option {option}: {text if text else '[Option text not detected]'}\n"
        return result
    def show_about_info(self):
        import tkinter.messagebox # Local import for safety
        # Ensure VERSION is defined, e.g., VERSION = "1.0.0" at the top of the file
        tkinter.messagebox.showinfo("About", f"LM Studio Batch Processor v{VERSION}\n\nCreated by Simon Andrews")

    def process_files(self, total_files):
        """Process all image files in the queue."""
        processed_count = 0
        # Ensure os and time are imported at the top of the file.
        # Ensure self.config, self.vision_var, self.use_ocr_var, self.results_queue, 
        # self.processing_queue, self.ocr_results, self.stop_processing are initialized.

        endpoint = self.config['API'].get('endpoint', 'http://localhost:1234/v1/chat/completions')
        model = self.config['API'].get('model', 'gemma-3-12b-instruct')
        has_vision = self.vision_var.get()
        use_ocr = self.use_ocr_var.get()
        
        try:
            temperature = float(self.config['API'].get('temperature', '0.7'))
        except ValueError:
            temperature = 0.7
            
        try:
            max_tokens = int(self.config['API'].get('max_tokens', '-1'))
        except ValueError:
            max_tokens = -1
            
        try:
            timeout = int(self.config['API'].get('timeout', '300'))
        except ValueError:
            timeout = 300
            
        try:
            max_retries = int(self.config['API'].get('max_retries', '2'))
        except ValueError:
            max_retries = 2
        
        while not self.processing_queue.empty() and not self.stop_processing:
            try:
                file_path = self.processing_queue.get_nowait()
                file_name = os.path.basename(file_path)

                self.results_queue.put(("status", f"Processing {file_name}..."))

                content = ""
                ocr_text = ""

                if use_ocr:
                    ocr_text = self.ocr_results.get(file_path, "")
                    if not ocr_text and getattr(self, 'ocr_thread_pool', None):  # Check if threaded OCR might be running
                        await_time = 0.0
                        max_ocr_wait = 2.0  # seconds
                        sleep_interval = 0.2
                        while not ocr_text and await_time < max_ocr_wait and not self.stop_processing:
                            time.sleep(sleep_interval)
                            await_time += sleep_interval
                            ocr_text = self.ocr_results.get(file_path, "")

                    if ocr_text:
                        # Ensure format_question_options handles potential errors gracefully
                        try:
                            formatted_ocr = self.format_question_options(ocr_text)
                            prompt = f"Analyze the following OCR text from an image named '{file_name}'. Extract the main question and its multiple-choice options (A, B, C, D). Format the output clearly.\n\nOCR Text:\n{formatted_ocr}"
                            content = prompt
                        except Exception as fmt_e:
                            self.results_queue.put(("error", f"Error formatting OCR for {file_name}: {str(fmt_e)}"))
                            content = f"Error formatting OCR for '{file_name}'. Based on filename, provide insights."
                    else:
                        content = f"OCR for '{file_name}' did not yield text or was not ready. Based on the filename, can you provide insights?"
                elif has_vision and not use_ocr:
                    try:
                        # Ensure encode_image handles potential errors gracefully
                        image_data = self.encode_image(file_path)
                        prompt = f"This image is named '{file_name}'. Analyze the image. If it contains a question with multiple-choice options (typically A, B, C, D), please extract the full question and all its options accurately. If not, describe the image content."
                        content = f"Here is an image to analyze:\n![Image]({image_data})\n\n{prompt}"
                    except Exception as enc_e:
                        self.results_queue.put(("error", f"Error encoding image {file_name}: {str(enc_e)}"))
                        processed_count += 1
                        self.results_queue.put(("progress", processed_count))
                        if hasattr(self.processing_queue, 'task_done'):
                            self.processing_queue.task_done()
                        continue  # to next file
                else:
                    # Fallback if OCR failed, in progress, or not used, and vision is not applicable
                    content = f"This is a prompt about an image file named '{file_name}'. Based on the filename, can you provide any insights or information that might be relevant?"

                # Make request to LM Studio API with retry logic
                response = None
                retries = 0
                last_error = None

                while response is None and retries <= max_retries:
                    if retries > 0:
                        retry_delay = retries * 2  # Progressive backoff
                        self.results_queue.put(("status", f"Retry {retries}/{max_retries} for {file_name} (waiting {retry_delay}s)..."))
                        time.sleep(retry_delay)

                    try:
                        response = self.get_lmstudio_response(
                            endpoint,
                            model,
                            content,
                            temperature,
                            max_tokens,
                            timeout
                        )
                    except Exception as e:
                        last_error = str(e)
                        self.results_queue.put(("status", f"Error: {last_error}. Retrying..."))
                        retries += 1

                if response:
                    # Store the API result now, we'll merge with final OCR result when creating document
                    self.results_queue.put(("api_result", (file_path, response)))
                    processed_count += 1
                    self.results_queue.put(("progress", processed_count))
                    self.results_queue.put(("status", f"Processed {processed_count}/{total_files}: {file_name}"))
                else:
                    # Report API error for this file
                    error_msg = f"API Error for {file_name}"
                    if last_error:
                        error_msg += f": {last_error}"
                    self.results_queue.put(("error", error_msg))
                    processed_count += 1
                    self.results_queue.put(("progress", processed_count))

                # Small delay to avoid overwhelming the API
                time.sleep(1)
            except Exception as e:
                self.results_queue.put(("error", f"Error processing {os.path.basename(file_path)}: {str(e)}"))
                processed_count += 1
                self.results_queue.put(("progress", processed_count))


        
        # Processing complete
        if not self.stop_processing:
            self.results_queue.put(("status", "Finalizing results..."))
            # Wait for any remaining OCR tasks to complete
            if self.ocr_thread_pool:
                self.ocr_thread_pool.shutdown(wait=True)
            self.results_queue.put(("complete", f"Processing complete. Processed {processed_count} files successfully."))
    
    def encode_image(self, image_path):
        """Encode image to base64 for API request."""
        # Open image and resize if too large
        try:
            img = Image.open(image_path)
            
            # Resize large images to save bandwidth
            max_size = 1024
            if img.width > max_size or img.height > max_size:
                img.thumbnail((max_size, max_size), Image.LANCZOS)
            
            # Convert to JPEG format
            if img.format != 'JPEG':
                img_buffer = io.BytesIO()
                img.convert('RGB').save(img_buffer, format='JPEG')
                img_buffer.seek(0)
                img_data = img_buffer.read()
            else:
                with open(image_path, "rb") as image_file:
                    img_data = image_file.read()
            
            # Encode to base64
            base64_encoded = base64.b64encode(img_data).decode('utf-8')
            return f"data:image/jpeg;base64,{base64_encoded}"
        
        except Exception as e:
            raise Exception(f"Error encoding image: {str(e)}")
    
    def get_lmstudio_response(self, endpoint, model, content, temperature, max_tokens, timeout=300):
        """Send a request to the LM Studio API and get a response."""
        headers = {
            "Content-Type": "application/json"
        }
        
        data = {
            "model": model,
            "messages": [
                {"role": "user", "content": content}
            ],
            "temperature": temperature,
            "stream": False
        }
        
        if max_tokens > 0:
            data["max_tokens"] = max_tokens
        
        response = requests.post(endpoint, headers=headers, json=data, timeout=timeout)
        response.raise_for_status()  # Raise exception for HTTP errors
        
        result = response.json()
        if "choices" in result and len(result["choices"]) > 0:
            return result["choices"][0]["message"]["content"]
        else:
            return None
    
    def extract_question_number(self, filename):
        """Extracts the question number from filename like name_p1240.txt -> 12"""
        # Match _p followed by digits, ignore trailing non-digits
        match = re.search(r'_p(\d+)\D*$', filename)
        if match:
            num_str = match.group(1)
            if len(num_str) > 2:
                 # Assume last two digits are not part of page number
                 return int(num_str[:-2])
            elif len(num_str) > 0:
                 return int(num_str)
        return float('inf')  # Return infinity for sorting if no number found
    
    def create_combined_word_document(self):
        """Create a combined Word document with all processed results."""
        if not self.processed_results:
            messagebox.showerror("Error", "No results to save.")
            return
        
        # Build final results with OCR data
        final_results = []
        for file_path, response in self.processed_results:
            # Get OCR text if available
            ocr_text = self.ocr_results.get(file_path, "")
            final_results.append((file_path, response, ocr_text))
        
        # Sort results by question number (if available in filename)
        final_results.sort(key=lambda x: self.extract_question_number(os.path.basename(x[0])))
        
        # Create a new document
        doc = Document()
        
        # Add a title
        title = doc.add_heading(f'LM Studio Batch Processing Results v{VERSION}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a timestamp
        doc.add_paragraph(f'Generated on: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')  # Add some space
        
        # Add each file and its response
        for result in final_results:
            file_path = result[0]
            response = result[1]
            ocr_text = result[2] if len(result) > 2 else ""
            
            filename = os.path.basename(file_path)
            
            # Add file name as heading
            doc.add_heading(filename, level=1)
            
            # Add the image
            try:
                doc.add_heading('Image:', level=2)
                doc.add_picture(file_path, width=Inches(5))
            except Exception as e:
                doc.add_paragraph(f"Could not add image: {str(e)}")
            
            # Add OCR text if available
            if ocr_text:
                doc.add_heading('OCR Text:', level=2)
                doc.add_paragraph(ocr_text)
            
            # Add response
            doc.add_heading('Response:', level=2)
            doc.add_paragraph(response)
            
            # Add separator
            doc.add_paragraph('-----------------------------------')
        
        # Ask for save location
        initial_dir = self.config['Settings'].get('last_folder', os.path.expanduser("~"))
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialdir=initial_dir,
            title="Save Combined Document As"
        )
        
        if save_path:
            try:
                doc.save(save_path)
                messagebox.showinfo(f"Success - v{VERSION}", "Document created successfully.")
            except Exception as e:
                messagebox.showerror(f"Error - v{VERSION}", f"Error saving document: {str(e)}")
    
    def check_progress(self):
        """Check the results queue for updates."""
        try:
            while not self.results_queue.empty():
                update_type, data = self.results_queue.get_nowait()
                
                if update_type == "status":
                    self.progress_var.set(data)
                
                elif update_type == "progress":
                    self.progress_bar["value"] = data
                
                elif update_type == "api_result":
                    # Store API results (file_path, response)
                    file_path, response = data
                    self.processed_results.append((file_path, response))
                
                elif update_type == "ocr_result":
                    # OCR results already stored in self.ocr_results
                    pass
                
                elif update_type == "error":
                    print(f"Error: {data}")
                
                elif update_type == "complete":
                    self.progress_var.set(data)
                    self.process_button.config(state=tk.NORMAL)
                    self.stop_button.config(state=tk.DISABLED)
                    
                    if self.processed_results:
                        if messagebox.askyesno(f"Processing Complete - v{VERSION}", 
                                             "Processing complete. Create combined document?"):
                            self.create_combined_word_document()
                    else:
                        messagebox.showinfo(f"Processing Complete - v{VERSION}", "No successful results to save.")
                    
                    return  # Stop checking for updates
            
            # Continue checking if processing is not complete
            if self.processing_thread and self.processing_thread.is_alive():
                self.master.after(100, self.check_progress)
            else:
                self.process_button.config(state=tk.NORMAL)
                self.stop_button.config(state=tk.DISABLED)
                
        except Exception as e:
            messagebox.showerror("Error", f"Error in progress update: {str(e)}")
    
    def stop_processing_thread(self):
        """Stop the processing thread."""
        self.stop_processing = True
        self.progress_var.set("Stopping...")
        self.stop_button.config(state=tk.DISABLED)
        
        # Shutdown OCR thread pool
        if self.ocr_thread_pool:
            self.ocr_thread_pool.shutdown(wait=False)

    def test_ocr(self):
        """Test OCR functionality with a test image."""
        # Get the tesseract path from the UI
        tesseract_path = self.tesseract_var.get()
        
        # Check if any file is selected in the listbox
        selected_indices = self.file_listbox.curselection()
        test_file = None
        
        if selected_indices:
            # Use the first selected file for testing
            selected_file = self.file_listbox.get(selected_indices[0])
            folder_path = self.folder_var.get()
            test_file = os.path.join(folder_path, selected_file)
            
            # Save the selected file to config
            self.config['Settings']['last_selected_file'] = selected_file
            self.save_config()
        
        if not test_file or not os.path.exists(test_file):
            # No file selected or it doesn't exist
            messagebox.showinfo(f"Test OCR - v{VERSION}", "Please select an image file from the list to test OCR.")
            return
        
        # Try to set the tesseract command
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Attempt OCR and show result
        try:
            self.master.config(cursor="wait")
            self.master.update()
            
            # Get current enhancement settings
            current_enhancement = self.ocr_enhance_var.get()
            psm = self.psm_var.get()
            language = self.lang_var.get()
            
            # Create window for results
            test_window = tk.Toplevel(self.master)
            test_window.title(f"OCR Test Results - v{VERSION}")
            test_window.geometry("800x600")
            
            # Add a label with the filename
            tk.Label(test_window, text=f"OCR for: {os.path.basename(test_file)}").pack(pady=5)
            
            # Add preview link for macOS
            if platform.system() == "Darwin":
                preview_frame = tk.Frame(test_window)
                preview_frame.pack(pady=2)
                
                # Create a hyperlink-style button to open in Preview
                preview_link = tk.Label(
                    preview_frame, 
                    text="View original image in Preview", 
                    fg="blue", 
                    cursor="hand2"
                )
                preview_link.pack()
                
                # Add underline on hover
                def on_enter(e):
                    preview_link.config(font=('TkDefaultFont', 9, 'underline'))
                def on_leave(e):
                    preview_link.config(font=('TkDefaultFont', 9))
                
                preview_link.bind("<Enter>", on_enter)
                preview_link.bind("<Leave>", on_leave)
                
                # Add click handler to open in Preview
                def open_in_preview(event):
                    try:
                        subprocess.run(["open", "-a", "Preview", test_file], check=False)
                    except Exception as e:
                        print(f"Error opening Preview: {e}")
                
                preview_link.bind("<Button-1>", open_in_preview)
            
            # Create notebook for different enhancement levels
            notebook = ttk.Notebook(test_window)
            notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
            
            # Dictionary to store tabs for each enhancement level
            tabs = {}
            results = {}
            
            # Test all enhancement levels in a specific order
            levels_to_test = ["Heavy", "Medium", "Light", "Adaptive", "None"]
            
            for level in levels_to_test:
                # Set the enhancement level
                self.ocr_enhance_var.set(level)
                
                # Run OCR with this enhancement level
                level_result = self.perform_ocr(test_file)
                results[level] = level_result
                
                # Create a tab for this level
                tab = ttk.Frame(notebook)
                tabs[level] = tab
                notebook.add(tab, text=level)
                
                # Add results to the tab
                text_frame = tk.Frame(tab)
                text_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
                
                scrollbar = tk.Scrollbar(text_frame)
                scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
                
                text_widget = tk.Text(text_frame, wrap=tk.WORD, yscrollcommand=scrollbar.set)
                text_widget.pack(fill=tk.BOTH, expand=True)
                
                scrollbar.config(command=text_widget.yview)
                text_widget.insert(tk.END, level_result)
                
            # Add a special tab for the current settings if not already tested
            if current_enhancement not in levels_to_test:
                # Restore the original enhancement setting
                self.ocr_enhance_var.set(current_enhancement)
                result = self.perform_ocr(test_file)
                
                # Create a tab for the current settings
                current_tab = ttk.Frame(notebook)
                notebook.add(current_tab, text=f"Current Settings ({current_enhancement})")
                
                # Add results to the tab
                text_frame = tk.Frame(current_tab)
                text_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
                
                scrollbar = tk.Scrollbar(text_frame)
                scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
                
                text_widget = tk.Text(text_frame, wrap=tk.WORD, yscrollcommand=scrollbar.set)
                text_widget.pack(fill=tk.BOTH, expand=True)
                
                scrollbar.config(command=text_widget.yview)
                text_widget.insert(tk.END, result)
            else:
                # Restore original setting
                self.ocr_enhance_var.set(current_enhancement)
            
            # Select the Heavy tab by default
            notebook.select(0)  # First tab (Heavy)
            
            # Add a settings display
            settings_frame = tk.Frame(test_window)
            settings_frame.pack(fill=tk.X, pady=5)
            
            tk.Label(settings_frame, text=f"Current Settings - Enhancement: {current_enhancement}, Layout: {psm}, Language: {language}").pack()
            
            # Add buttons frame
            buttons_frame = tk.Frame(test_window)
            buttons_frame.pack(pady=10)
            
            # Add close button
            tk.Button(buttons_frame, text="Close", command=test_window.destroy).pack(side=tk.LEFT, padx=5)
            
            # Add apply settings button for Heavy if successful
            heavy_result = results.get("Heavy", "")
            if heavy_result and not heavy_result.startswith("OCR Error"):
                def apply_heavy_settings():
                    self.use_ocr_var.set(True)
                    self.ocr_enhance_var.set("Heavy")  # Set to Heavy
                    self.config['Settings']['use_ocr'] = 'True'
                    self.config['Settings']['tesseract_path'] = tesseract_path
                    self.save_config()
                    messagebox.showinfo(f"OCR Test - v{VERSION}", "OCR settings saved with Heavy enhancement and OCR has been enabled.")
                    test_window.destroy()
                    
                tk.Button(buttons_frame, text="Use Heavy Enhancement", command=apply_heavy_settings).pack(side=tk.LEFT, padx=5)
            
            # Also add button for current settings if different from Heavy
            if current_enhancement != "Heavy" and not (results.get(current_enhancement, "")).startswith("OCR Error"):
                def apply_current_settings():
                    self.use_ocr_var.set(True)
                    self.config['Settings']['use_ocr'] = 'True'
                    self.config['Settings']['tesseract_path'] = tesseract_path
                    self.save_config()
                    messagebox.showinfo(f"OCR Test - v{VERSION}", f"OCR settings saved with {current_enhancement} enhancement and OCR has been enabled.")
                    test_window.destroy()
                    
                tk.Button(buttons_frame, text=f"Use {current_enhancement} Enhancement", command=apply_current_settings).pack(side=tk.LEFT, padx=5)
            
        except Exception as e:
            messagebox.showerror("OCR Test Error", f"Error testing OCR: {str(e)}")
        finally:
            self.master.config(cursor="")

    def show_about_info(self):
        """Show information about the application."""
        about_text = f"""
LM Studio Batch Processor v{VERSION}

A tool for processing batches of images through LM Studio's API.
Features include:
- OCR text extraction with multiple enhancement levels
- Integration with OpenAI-compatible API interface
- Parallel processing for improved performance
- Comprehensive document generation

© {time.strftime('%Y')}
        """
        messagebox.showinfo(f"About - v{VERSION}", about_text)

    def quit_application(self):
        """Exit the application cleanly."""
        if self.processing_thread and self.processing_thread.is_alive():
            if messagebox.askyesno(f"Quit - v{VERSION}", "Processing is still running. Are you sure you want to quit?"):
                self.stop_processing_thread()
                self.master.after(500, self.master.destroy)
            return
        self.master.destroy()

def main():
    root = tk.Tk()
    app = BatchProcessor(root)
    
    # Maximize the window
    root.update()  # Update to get correct screen dimensions
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    root.geometry(f"{screen_width}x{screen_height}+0+0")
    
    # Alternative maximization methods based on platform
    try:
        # For Windows
        if platform.system() == "Windows":
            root.state('zoomed')
        # For macOS - use available screen space minus dock/menubar
        elif platform.system() == "Darwin":
            # Leave some space for menu bar and dock
            root.geometry(f"{screen_width}x{screen_height-80}+0+0")
            # Try to use zoomed state if available (might not work on all macOS versions)
            try:
                root.attributes('-zoomed', True)
            except:
                pass
    except Exception as e:
        print(f"Could not maximize window: {e}")
    
    # Make window active and bring to foreground
    root.lift()  # Bring window to top
    root.attributes('-topmost', True)  # Make it stay on top temporarily
    root.update()  # Force update
    root.attributes('-topmost', False)  # Remove always-on-top
    root.focus_force()  # Force focus to this window
    
    # For macOS - additional method to activate app
    if platform.system() == "Darwin":
        try:
            # Try using osascript to activate the application
            import subprocess
            app_name = "Python" if "Python" in root.winfo_name() else root.winfo_name()
            subprocess.run(["osascript", "-e", f'tell application "{app_name}" to activate'], 
                          check=False, capture_output=True)
        except Exception as e:
            print(f"Could not activate window via AppleScript: {e}")
    
    root.mainloop()

if __name__ == "__main__":
    main() 