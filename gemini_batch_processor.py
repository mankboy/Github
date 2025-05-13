import os
import time
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import queue
import configparser
import re
from gemini_image_analyzer import (
    load_config, save_config, get_api_key, 
    encode_image, get_final_response_rag
)
# Import docx components directly for combined document creation
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration file (shared with the main analyzer)
CONFIG_FILE = os.path.expanduser("~/.gemini_config.ini")

class BatchProcessor:
    def __init__(self, master):
        self.master = master
        master.title("Gemini Batch Image Processor")
        master.geometry("600x500")
        
        # Load configuration
        self.config = load_config()
        
        # Setup UI elements
        self.create_widgets()
        
        # Processing variables
        self.processing_queue = queue.Queue()
        self.results_queue = queue.Queue()
        self.processed_results = [] # Store successful (path, response) tuples
        self.processing_thread = None
        self.stop_processing = False
        
    def create_widgets(self):
        # Top frame for folder selection
        folder_frame = tk.Frame(self.master)
        folder_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(folder_frame, text="Folder:").pack(side=tk.LEFT)
        
        self.folder_var = tk.StringVar()
        folder_entry = tk.Entry(folder_frame, textvariable=self.folder_var, width=50)
        folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        browse_button = tk.Button(folder_frame, text="Browse", command=self.browse_folder)
        browse_button.pack(side=tk.LEFT)
        
        # Frame for file list
        files_frame = tk.Frame(self.master)
        files_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        tk.Label(files_frame, text="Images to process:").pack(anchor=tk.W)
        
        # Create a frame with scrollbars for the file list
        file_list_frame = tk.Frame(files_frame)
        file_list_frame.pack(fill=tk.BOTH, expand=True)
        
        # Add scrollbars to the file list
        scrollbar_y = tk.Scrollbar(file_list_frame)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        
        scrollbar_x = tk.Scrollbar(file_list_frame, orient=tk.HORIZONTAL)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.file_listbox = tk.Listbox(
            file_list_frame,
            selectmode=tk.EXTENDED,
            yscrollcommand=scrollbar_y.set,
            xscrollcommand=scrollbar_x.set
        )
        self.file_listbox.pack(fill=tk.BOTH, expand=True)
        
        scrollbar_y.config(command=self.file_listbox.yview)
        scrollbar_x.config(command=self.file_listbox.xview)
        
        # Control buttons
        control_frame = tk.Frame(self.master)
        control_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.process_button = tk.Button(
            control_frame, 
            text="Process & Create Combined Doc", 
            command=self.start_processing
        )
        self.process_button.pack(side=tk.LEFT, padx=5)
        
        self.stop_button = tk.Button(
            control_frame, 
            text="Stop", 
            command=self.stop_processing_thread,
            state=tk.DISABLED
        )
        self.stop_button.pack(side=tk.LEFT, padx=5)
        
        # Progress frame
        progress_frame = tk.Frame(self.master)
        progress_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(progress_frame, text="Progress:").pack(anchor=tk.W)
        
        self.progress_var = tk.StringVar(value="Ready")
        tk.Label(progress_frame, textvariable=self.progress_var).pack(anchor=tk.W)
        
        self.progress_bar = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=100, mode='determinate')
        self.progress_bar.pack(fill=tk.X, pady=5)
    
    def browse_folder(self):
        """Open a folder browser dialog."""
        initial_dir = self.config['Settings'].get('last_folder', os.path.expanduser("~"))
        
        # Ensure the initial directory exists
        if not os.path.exists(initial_dir):
            initial_dir = os.path.expanduser("~")
        
        folder_path = filedialog.askdirectory(
            title="Select folder containing images",
            initialdir=initial_dir
        )
        
        if folder_path:
            # Update last folder in config
            self.config['Settings']['last_folder'] = folder_path
            save_config(self.config)
            
            # Update the entry and list the files
            self.folder_var.set(folder_path)
            self.list_image_files(folder_path)
    
    def list_image_files(self, folder_path):
        """List all image files in the selected folder."""
        self.file_listbox.delete(0, tk.END)
        
        if not folder_path or not os.path.exists(folder_path):
            return
        
        for file in os.listdir(folder_path):
            if file.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
                self.file_listbox.insert(tk.END, file)
    
    def start_processing(self):
        """Start processing the images in a separate thread."""
        folder_path = self.folder_var.get()
        
        if not folder_path or not os.path.exists(folder_path):
            messagebox.showerror("Error", "Please select a valid folder.")
            return
        
        # Check if API key exists
        api_key = get_api_key(self.config)
        if not api_key:
            messagebox.showerror("Error", "API key is required.")
            return
        
        # Get all selected files, or all files if none selected
        selected_indices = self.file_listbox.curselection()
        if selected_indices:
            selected_files = [self.file_listbox.get(i) for i in selected_indices]
        else:
            selected_files = [self.file_listbox.get(i) for i in range(self.file_listbox.size())]
        
        if not selected_files:
            messagebox.showerror("Error", "No image files found in the selected folder.")
            return
        
        # Reset results list
        self.processed_results = []
        
        # Clear queues
        while not self.processing_queue.empty():
            self.processing_queue.get()
        
        while not self.results_queue.empty():
            self.results_queue.get()
        
        # Add files to processing queue
        for file in selected_files:
            file_path = os.path.join(folder_path, file)
            self.processing_queue.put(file_path)
        
        # Start processing thread
        self.stop_processing = False
        self.processing_thread = threading.Thread(target=self.process_images, args=(api_key, len(selected_files)))
        self.processing_thread.daemon = True
        self.processing_thread.start()
        
        # Update UI
        self.process_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)
        self.progress_var.set("Processing...")
        self.progress_bar["value"] = 0
        self.progress_bar["maximum"] = len(selected_files)
        
        # Start checking for results
        self.master.after(100, self.check_progress)
    
    def process_images(self, api_key, total_files):
        """Process all images in the queue."""
        processed_count = 0
        
        while not self.processing_queue.empty() and not self.stop_processing:
            try:
                file_path = self.processing_queue.get()
                
                # Update status
                self.results_queue.put(("status", f"Processing {os.path.basename(file_path)}..."))
                
                # Get response from Gemini using RAG function (passing empty search results)
                gemini_response = get_final_response_rag(api_key, file_path, "")
                
                if gemini_response:
                    # Store successful result instead of creating document
                    self.results_queue.put(("result", (file_path, gemini_response)))
                    processed_count += 1
                    self.results_queue.put(("progress", processed_count))
                    self.results_queue.put(("status", f"Processed {processed_count}/{total_files}: {os.path.basename(file_path)}"))
                else:
                    # Report API error for this file
                    self.results_queue.put(("error", f"API Error for {os.path.basename(file_path)}"))
                    # Update progress bar even if error, so it reaches max eventually
                    # self.results_queue.put(("progress", processed_count + (total_files - self.processing_queue.qsize() - processed_count))) # This logic is complex, maybe just increment progress?
                
                # Small delay to avoid overwhelming the API
                time.sleep(2)  # Increased delay for Gemini API rate limiting
                
            except Exception as e:
                self.results_queue.put(("error", f"Error processing {os.path.basename(file_path)}: {str(e)}"))
        
        # Processing complete
        if not self.stop_processing:
            self.results_queue.put(("complete", f"Processing complete. Processed {processed_count} files successfully."))
    
    def extract_question_number(self, filename):
        """Extracts the question number from filename like name_p1240.png -> 12"""
        # Match _p followed by digits, ignore trailing non-digits
        match = re.search(r'_p(\d+)\D*$', filename)
        if match:
            num_str = match.group(1)
            if len(num_str) > 2:
                 # Assume last two digits are not part of page number based on example _p1240 -> 12
                 # This assumption might need adjustment if patterns vary
                 return int(num_str[:-2])
            elif len(num_str) > 0:
                 return int(num_str) # Handle cases like _p5 or _p15
        return float('inf') # Return infinity for sorting if no number found
    
    def create_combined_word_document(self):
        """Creates a single Word document from all processed results."""
        if not self.processed_results:
            messagebox.showinfo("Info", "No results to save.")
            return
        
        # Sort results based on extracted question number
        try:
            self.processed_results.sort(key=lambda item: self.extract_question_number(os.path.basename(item[0])))
            total_questions = len(self.processed_results)
        except Exception as e:
            print(f"Error sorting results: {e}")
            messagebox.showerror("Error", "Could not sort results based on filenames.")
            # Proceed without sorting if error occurs?
            total_questions = len(self.processed_results)
        
        # Ask user for save location
        initial_dir = self.config['Settings'].get('last_folder', os.path.expanduser("~"))
        output_path = filedialog.asksaveasfilename(
            title="Save Combined Analysis As",
            initialdir=initial_dir,
            defaultextension=".docx",
            initialfile="Combined_Analysis.docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")]
        )
        
        if not output_path:
            print("Save cancelled.")
            return
        
        try:
            doc = Document()
            doc.add_heading('Combined Gemini Analysis', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, (image_path, gemini_response) in enumerate(self.processed_results):
                question_num = self.extract_question_number(os.path.basename(image_path))
                question_num_display = str(question_num) if question_num != float('inf') else f"(Unknown {i+1})"
                
                doc.add_heading(f'Question {question_num_display} of {total_questions}', level=1)
                
                doc.add_paragraph(f"Source Image: {os.path.basename(image_path)}")
                try:
                    max_width = Inches(6.0)
                    doc.add_picture(image_path, width=max_width)
                except Exception as e:
                    doc.add_paragraph(f"[Error including image: {str(e)}]")
                
                doc.add_paragraph()
                
                doc.add_heading('Gemini Analysis:', level=2)
                try:
                    content = gemini_response["choices"][0]["message"]["content"]
                    
                    # --- Debugging Print --- 
                    print(f"\n--- Raw Content for Question {question_num_display} ({os.path.basename(image_path)}) ---\n")
                    print(content)
                    print("---------------------------------------------------\n")
                    # --- End Debugging Print ---

                    current_heading = None
                    for line in content.split('\n'):
                        line = line.strip()
                        if not line:
                            continue
                        if line.startswith('# '):
                            current_heading = line[2:]
                            doc.add_heading(current_heading, level=2) # Treat as level 2 under question
                        elif line.startswith('## '):
                            current_heading = line[3:]
                            doc.add_heading(current_heading, level=3)
                        else:
                            p = doc.add_paragraph(line)
                            if current_heading and ("Why other answers" in current_heading):
                                p.paragraph_format.left_indent = Inches(0.25)
                except Exception as e:
                    doc.add_paragraph(f"Error formatting Gemini response: {str(e)}")
                    doc.add_paragraph(str(gemini_response))
                
                # Add page break unless it's the last item
                if i < len(self.processed_results) - 1:
                    doc.add_page_break()
            
            doc.save(output_path)
            messagebox.showinfo("Success", f"Combined analysis saved to:\n{output_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create combined document: {str(e)}")
            print(f"Error creating combined document: {str(e)}")
    
    def check_progress(self):
        """Check for updates from the processing thread."""
        try:
            while True:
                message_type, message = self.results_queue.get_nowait()
                
                if message_type == "status":
                    self.progress_var.set(message)
                elif message_type == "progress":
                    self.progress_bar["value"] = message
                elif message_type == "result":
                    # Store successful results
                    self.processed_results.append(message)
                elif message_type == "error":
                    print(f"Error: {message}") # Log errors
                    # Optionally update progress bar here if needed
                elif message_type == "complete":
                    self.progress_var.set(message)
                    self.process_button.config(state=tk.NORMAL)
                    self.stop_button.config(state=tk.DISABLED)
                    # Create the combined document *after* processing is complete
                    self.create_combined_word_document()
                    return # Exit the loop
        except queue.Empty:
            # If nothing in the queue, check again later
            if self.processing_thread and self.processing_thread.is_alive():
                self.master.after(100, self.check_progress)
            else:
                # Thread finished unexpectedly or was stopped
                final_status = "Processing stopped." if self.stop_processing else "Processing finished unexpectedly."
                if not self.stop_processing and self.processed_results:
                    # If processing finished without 'complete' but we have results, try saving.
                     final_status = "Processing finished. Creating document..."
                     self.progress_var.set(final_status)
                     self.create_combined_word_document()
                else:
                     self.progress_var.set(final_status)
                
                self.process_button.config(state=tk.NORMAL)
                self.stop_button.config(state=tk.DISABLED)
    
    def stop_processing_thread(self):
        """Stop the processing thread."""
        self.stop_processing = True
        self.stop_button.config(state=tk.DISABLED)
        self.progress_var.set("Stopping...")

def main():
    root = tk.Tk()
    app = BatchProcessor(root)
    root.lift() # Bring window to front
    root.focus_force() # Force focus to the window
    root.mainloop()

if __name__ == "__main__":
    main() 